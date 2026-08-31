#!/usr/bin/env python3
"""
Generador Maestro: Guía 1 - Definición Proyecto APT (Fase 1)
Asignatura: Portafolio de Título (APT122) - Duoc UC Sede San Joaquín
Carrera: Ingeniería en Conectividad y Redes
Estudiante: Bruno Urrea Ortiz (RUT: 21.543.637-3)

Genera las 3 versiones oficiales:
1. Markdown Oficial (.md)
2. Word Oficial con Portada e Índice (.docx)
3. PDF Oficial Institucional (.pdf)
"""

import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether, HRFlowable
)
from reportlab.pdfgen import canvas

OUTPUT_DIR = "docs/Fase_1_Definicion_Proyecto_APT"
MD_FILE = os.path.join(OUTPUT_DIR, "Guia1_Definicion_Proyecto_APT_Fase1_Bruno_Urrea.md")
DOCX_FILE = os.path.join(OUTPUT_DIR, "Guia1_Definicion_Proyecto_APT_Fase1_Bruno_Urrea.docx")
PDF_FILE = os.path.join(OUTPUT_DIR, "Guia1_Definicion_Proyecto_APT_Fase1_Bruno_Urrea.pdf")

# =============================================================================
# 1. CONTENIDO COMPLETO EN MARKDOWN (ESTRUCTURA OFICIAL DUOC UC)
# =============================================================================
MARKDOWN_CONTENT = """# Guía 1. Definición Proyecto APT
### Asignatura: Portafolio de Título (APT122) — Asignatura Capstone (Fase 1)
**Institución:** Duoc UC — Sede San Joaquín  
**Escuela:** Escuela de Informática y Telecomunicaciones  
**Carrera:** Ingeniería en Conectividad y Redes  
**Estudiante:** Bruno Urrea Ortiz  
**RUT:** 21.543.637-3  
**Docente Guía:** Comisión Evaluadora APT122  
**Fecha:** Septiembre 2026  

---

## Índice de Contenidos
1. [Portada Institucional](#portada-institucional)
2. [A. PARTE I](#a-parte-i)
   - [1. Antecedentes Personales](#1-antecedentes-personales)
   - [2. Descripción Proyecto APT](#2-descripción-proyecto-apt)
   - [3. Fundamentación Proyecto APT](#3-fundamentación-proyecto-apt)
     - [3.1 Relevancia del proyecto APT](#31-relevancia-del-proyecto-apt)
     - [3.2 Descripción del Proyecto APT](#32-descripción-del-proyecto-apt)
     - [3.3 Pertinencia del proyecto con el perfil de egreso](#33-pertinencia-del-proyecto-con-el-perfil-de-egreso)
     - [3.4 Relación con los intereses profesionales](#34-relación-con-los-intereses-profesionales)
     - [3.5 Factibilidad de desarrollo del Proyecto APT](#35-factibilidad-de-desarrollo-del-proyecto-apt)
3. [B. PARTE II](#b-parte-ii)
   - [4. Objetivos](#4-objetivos)
     - [4.1 Objetivo General](#41-objetivo-general)
     - [4.2 Objetivos Específicos (Formulados como Acciones)](#42-objetivos-específicos-formulados-como-acciones)
   - [5. Metodología](#5-metodología)
     - [5.1 Descripción de la Metodología](#51-descripción-de-la-metodología)
     - [5.2 Funciones, Tareas y Responsabilidades del Equipo](#52-funciones-tareas-y-responsabilidades-del-equipo)
   - [6. Evidencias (Tabla Oficial)](#6-evidencias)
   - [7. Plan de Trabajo (Tabla Oficial con Recursos y Factores)](#7-plan-de-trabajo)
   - [8. Carta Gantt (18 Semanas Académicas)](#8-carta-gantt)

---

# A. PARTE I

## 1. Antecedentes Personales
A continuación, se presenta la tabla en la que se completa la información solicitada:

| Campo | Información Solicitada |
| :--- | :--- |
| **Nombre estudiante** | Bruno Urrea Ortiz |
| **Rut** | 21.543.637-3 |
| **Carrera** | Ingeniería en Conectividad y Redes |
| **Sede** | San Joaquín |
| **Integrantes del Equipo (Grupal)** | • **Bruno Urrea Ortiz:** Líder de Ciberseguridad, Motor KRONOS y Gemini Live API.<br>• **Freddy Vásquez Cortés:** Ingeniería de Routing, Switching perimetral y Telefonía Asterisk PBX.<br>• **Cristóbal Quezada:** Administración de Servicios Web, Proxy Inverso HAProxy y DMZ.<br>• **Kevin Retamales:** Hardening Perimetral, Inteligencia pfBlockerNG y Control de Calidad QA. |

---

## 2. Descripción Proyecto APT
En la descripción se señala brevemente el nombre del proyecto APT, las áreas de desempeño y las competencias del perfil de egreso que se ponen en práctica:

| Campo | Detalle |
| :--- | :--- |
| **Nombre del proyecto** | **KRONOS SENTINEL:** Autonomous AI-IPS & Real-Time Incident Voice Response SOAR Architecture |
| **Área(s) de desempeño** | • **Seguridad de Sistemas y Redes:** Ciberseguridad Defensiva, Hardening Perimetral y Prevención de Intrusiones (IPS).<br>• **Administración de Redes y Telecomunicaciones:** Infraestructura Corporativa L2/L3 y Segmentación en VLANs 802.1Q.<br>• **Comunicación Unificada y Corporativa:** Telefonía IP Asterisk (PJSIP), Dialplans y Streaming de Voz en Tiempo Real.<br>• **Análisis de Soluciones de Conectividad:** Automatización en Python 3.12, Orquestación SOAR e Inteligencia Artificial Multimodal. |
| **Competencias** | • **Competencia 8:** Crear planes de prevención y respuesta a riesgos informáticos en la red de acuerdo con estándares, normativas y buenas prácticas de la industria.<br>• **Competencia 6:** Automatizar procesos y gestión de plataformas de red mediante programación y scripting avanzado.<br>• **Competencia 5:** Unificar servicios de voz, datos y video en redes de clientes asegurando calidad de servicio (QoS) y convergencia tecnológica.<br>• **Competencia 4:** Controlar y operar redes corporativas de gran tamaño manteniendo la continuidad operativa y alta disponibilidad.<br>• **Competencia 3:** Adaptar tecnologías de punta y tendencias tecnológicas emergentes para resolver necesidades de innovación organizacional.<br>• **Competencia 7:** Gestionar la seguridad de la información frente a vulnerabilidades en aplicaciones y servicios de red. |

---

## 3. Fundamentación Proyecto APT

| Campo | Detalle y Justificación |
| :--- | :--- |
| **Relevancia del proyecto APT** | **¿Por qué se escogió este tema y cuál es su relevancia laboral?**<br>En los Centros de Operaciones de Seguridad (SOC) modernos existen dos fallas críticas:<br>1. *Crisis de hiper-alerta y falsos positivos (>50%):* Los motores de inspección profunda tradicionales (Suricata/Snort) saturan al analista con alertas inocuas.<br>2. *Colapso cognitivo del operador bajo ataque crítico:* Ante una intrusión grave (SQLi/RCE), el operador entra en visión de túnel al intentar aislar el firewall por CLI y redactar reportes al CISO simultáneamente, provocando demoras de horas y errores de tipeo.<br><br>**¿Dónde se ubica la situación?**<br>En datacenters corporativos, entidades financieras, plataformas web expuestas e infraestructuras empresariales de Chile y Latinoamérica.<br><br>**¿A quiénes afecta o impacta?**<br>Impacta a Analistas SOC N1/N2/N3, CISOs, Directores de TI y a la continuidad del negocio corporativo.<br><br>**¿Cuál es el aporte de valor?**<br>KRONOS SENTINEL aporta una solución SOAR de Costo Cero ($0 CLP) que desacopla la contención técnica atómica en kernel de FreeBSD (<100 ms) de la comunicación estratégica mediante llamadas telefónicas con Inteligencia Artificial conversacional (Gemini Live Flash 3.1) en tiempo real (<1.5 s). |
| **Descripción del Proyecto APT** | **Propósito Central y Enfoque Metodológico:**<br>Diseñar, implementar y validar una arquitectura de defensa en profundidad de 4 pilares tecnológicos:<br>1. *Perímetro Fortificado:* Firewall **pfSense CE 2.9.0** con **Suricata 7.x Inline IPS (Netmap)** para descarte de paquetes en hardware sin latencia, más **pfBlockerNG-devel** con inteligencia GeoIP MaxMind y listas FireHOL.<br>2. *Capa Web DMZ:* Proxy inverso **HAProxy 2.8+** con terminación SSL/TLS (puerto 443) y *Stick-Tables* anti-fuzzing L7 protegiendo el entorno vulnerable controlado (**DVWA** en Docker, VLAN 20).<br>3. *Motor de Correlación KRONOS (Python 3.12):* Daemon asíncrono que ingesta `eve.json`, suprime >50% de falsos positivos mediante analizador sintáctico AST y ejecuta la herramienta de kernel de FreeBSD **`pfctl`** para terminar estados hostiles (`pfctl -k`) y verificar la tabla `<snort2c>`.<br>4. *Respuesta Telefónica Autónoma SOAR:* Centralita **Asterisk 20 LTS** en Docker con módulo AMI (`Originate`) que timbra al softphone del CISO y conecta el flujo de audio RTP con **Google Gemini Live Flash 3.1** mediante WebSocket bidireccional (PCM 24kHz). |
| **Pertinencia del proyecto con el perfil de egreso** | El proyecto articula de forma rigurosa las 4 áreas troncales de la carrera:<br>• **Networking L2/L3:** Segmentación en 4 VLANs 802.1Q (Corp 10, DMZ 20, VoIP 30, Mgmt 99), direccionamiento CIDR y conmutación.<br>• **Seguridad Perimetral:** Hardening de kernel Unix/FreeBSD, reglas Zero Trust, inspección profunda Netmap e inteligencia de amenazas.<br>• **Telefonía VoIP:** Despliegue de centralita Asterisk PBX, protocolo SIP/PJSIP, códecs de audio y enlaces telefónicos en tiempo real.<br>• **Automatización y Programación:** Desarrollo de software en Python, manejo de sockets IPC, parsing estructurado JSON y consumo de APIs de IA Generativa. |
| **Relación con los intereses profesionales** | **¿Cuáles son tus intereses profesionales?**<br>Especialización en Ciberseguridad Defensiva (Blue Team), Arquitectura de Seguridad Perimetral, Hardening de Sistemas Unix/FreeBSD y Sistemas Autónomos de Orquestación y Respuesta ante Incidentes (SOAR).<br><br>**¿Qué aspectos se ven reflejados en el Proyecto APT?**<br>El diseño de infraestructuras Zero Trust, la mitigación atómica de ataques en kernel, el análisis heurístico de tráfico malicioso y la integración de IA generativa aplicada a operaciones de seguridad.<br><br>**¿Cómo contribuye a tu desarrollo profesional?**<br>Permite proyectar mi inserción laboral hacia cargos como Ingeniero de Ciberseguridad / Arquitecto SecOps en SOC corporativos o liderar equipos de respuesta ante amenazas APT, acreditando dominio práctico de tecnologías Enterprise de costo cero ($0 CLP) y alineación con certificaciones internacionales (Cisco CCNA, CompTIA Security+, NIST SP 800-207). |
| **Factibilidad de desarrollo del Proyecto APT** | **Acotación del Alcance a 18 Semanas con el Equipo Disponible:**<br>• *(1) Duración del semestre:* 18 semanas académicas cronogramadas en 4 fases de trabajo.<br>• *(2) Horas asignadas:* 72 horas presenciales de taller de titulación + 144 horas de laboratorio autónomo = 216 horas por integrante (864 horas de ingeniería grupal total).<br>• *(3) Materiales requeridos ($0 CLP):* 100% Open Source y capas comunitarias gratuitas (FreeBSD 14, pfSense CE 2.9.0, Suricata 7.x, HAProxy 2.8, Asterisk 20 LTS, Docker, Python 3.12, Google AI Studio Free Tier, MaxMind GeoLite2). Cero costo de licenciamiento comercial.<br>• *(4) Factores que facilitan el desarrollo:* Plataforma de virtualización Proxmox VE bare-metal / VMware, laboratorios de redes Duoc UC San Joaquín y modularidad de la arquitectura.<br>• *(5) Factores que dificultan y soluciones:*<br>   - *Dificultad:* Bloqueo de puertos o CGNAT en redes institucionales/remotas.<br>   - *Solución:* Malla VPN Zero Trust con **Tailscale Subnet Router** (túnel WireGuard) para acceder a la subred VoIP `192.168.30.0/24` sin requerir IP pública fija.<br>   - *Dificultad:* Saturación de llamadas por falsos positivos.<br>   - *Solución:* Filtro heurístico AST que exige confirmación atómica en la tabla `<snort2c>` antes de activar el auto-dialer AMI.<br>   - *Dificultad:* Latencia en respuesta de voz.<br>   - *Solución:* Conexión WebSocket directa PCM 24kHz con Gemini Live Flash 3.1 (<400 ms). |

---

# B. PARTE II

## 4. Objetivos

### 4.1 Objetivo General
Diseñar, implementar y validar una arquitectura de defensa en profundidad y respuesta autónoma ante incidentes (SOAR) de costo cero ($0 CLP) denominada **KRONOS SENTINEL**, integrando prevención de intrusiones en kernel, supresión heurística de falsos positivos e interlocución telefónica interactiva por voz en tiempo real con Inteligencia Artificial hacia los responsables de seguridad (CISO).

### 4.2 Objetivos Específicos (Formulados como Acciones)
1. **Diseñar** la topología lógica y física de red perimetral y el esquema de direccionamiento CIDR segmentado en 4 VLANs 802.1Q (Corporativa, DMZ, VoIP y Gestión) en pfSense CE 2.9.0.
2. **Implementar** la configuración del cortafuegos perimetral pfSense CE 2.9.0 con políticas Zero Trust, reglas de aislamiento entre subredes y optimización de hardware offloading para Netmap.
3. **Configurar** el motor de detección y prevención de intrusiones Suricata 7.x en modo Inline IPS sobre framework Netmap con firmas ET Open y reglas automáticas `dropsid.conf` para el descarte de paquetes maliciosos en hardware ring-buffer.
4. **Implementar** el filtrado geográfico y reputacional de amenazas con pfBlockerNG-devel utilizando bases de datos MaxMind GeoLite2 y listas de bloqueo global FireHOL L1 y Spamhaus DROP.
5. **Configurar** el proxy inverso HAProxy 2.8+ con terminación segura SSL/TLS (puerto 443) y *Stick-Tables* de control de tasa L7 en memoria RAM, protegiendo y publicando el servidor web de pruebas vulnerables DVWA en la DMZ.
6. **Desarrollar** en Python 3.12 el Motor de Correlación KRONOS con analizador estructurado de logs `eve.json` y algoritmo heurístico AST para suprimir más del 50% de falsos positivos en vectores de inyección SQL (SQLi) y ejecución remota de comandos (RCE).
7. **Integrar** el control atómico de estados y tablas en el kernel de FreeBSD mediante la herramienta de sistema `pfctl`, ejecutando la terminación inmediata de sesiones activas (`pfctl -k`) y verificando la persistencia del atacante en la tabla dinámica `<snort2c>`.
8. **Desplegar** la centralita de telefonía IP Asterisk 20 LTS en contenedor Docker con canal PJSIP, configuración de Dialplans y módulo auto-dialer AMI (`Originate`) para el timbrado prioritario al softphone del CISO.
9. **Construir** el puente de audio bidireccional por WebSocket hacia Google Gemini Live Flash 3.1 (PCM 24kHz), permitiendo la entrega de un debriefing táctico hablado en lenguaje natural y la recepción de instrucciones de respuesta por voz.
10. **Validar** el rendimiento, la latencia total de respuesta (<1.5 s) y la efectividad de la arquitectura completa mediante la ejecución de matrices de pruebas de aseguramiento de calidad (QA) y simulación de ataques en vivo.

---

## 5. Metodología

### 5.1 Descripción de la Metodología
Se aplicará una **Metodología de Ingeniería en Ciclo Iterativo e Incremental**, estructurada en 4 etapas operativas con control de calidad y pruebas de regresión continuas:
* **Etapa 1: Levantamiento, Diseño y Topología de Red (Semanas 1 a 4):** Modelamiento de direccionamiento IP, matriz de VLANs 802.1Q, hardening base del sistema operativo FreeBSD y configuración inicial de pfSense CE 2.9.0.
* **Etapa 2: Implementación de Perímetro y Capa Web DMZ (Semanas 5 a 10):** Despliegue de Suricata Inline Netmap IPS, configuración de pfBlockerNG GeoIP, parametrización de HAProxy SSL con Stick-Tables y puesta en marcha del laboratorio DVWA en Docker.
* **Etapa 3: Desarrollo de KRONOS, Telefonía PBX y Voz IA (Semanas 11 a 15):** Programación en Python 3.12 del analizador AST, creación de wrappers `pfctl`, construcción de imagen Docker Asterisk 20 LTS, desarrollo del auto-dialer AMI y bridge WebSocket de Gemini Live API.
* **Etapa 4: Pruebas de Integración QA, Documentación y Defensa (Semanas 16 a 18):** Pruebas de penetración SQLi en vivo, medición de tiempos de respuesta (<1.5 s), generación de manuales técnicos PDF y preparación de la presentación ejecutiva ante la comisión.

### 5.2 Funciones, Tareas y Responsabilidades del Equipo
* **Bruno Urrea Ortiz (Líder de Ciberseguridad, Motor KRONOS y Gemini Live):**
  - Dirección arquitectónica y diseño de la solución SOAR.
  - Desarrollo del Motor de Correlación KRONOS (`log_correlator.py`) y del filtro sintáctico AST (`false_positive_filter.py`).
  - Implementación de wrappers de kernel FreeBSD con `pfctl` (`pfctl_wrapper.py`).
  - Integración del cliente WebSocket seguro hacia Google Gemini Live Flash 3.1 (`gemini_live_client.py`).
* **Freddy Vásquez Cortés (Ingeniero de Routing, Switching y Telefonía Asterisk PBX):**
  - Configuración de la troncal 802.1Q y direccionamiento de subredes VLAN en pfSense.
  - Despliegue y hardening de la centralita Asterisk 20 LTS en Docker (`pjsip.conf`, `extensions.conf`).
  - Desarrollo del script auto-dialer AMI (`call_trigger.py`).
  - Configuración del enrutamiento VPN Zero Trust con Tailscale Subnet Router.
* **Cristóbal Quezada (Administrador de Servicios Web, Proxy HAProxy y DMZ):**
  - Configuración del frontend y backend de HAProxy 2.8+ con terminación HTTPS 443.
  - Parametrización de *Stick-Tables* de rate-limiting anti-fuzzing L7 en memoria RAM.
  - Despliegue, aislamiento y administración del contenedor vulnerable DVWA en VLAN 20 DMZ.
  - Hardening de protocolos TLS y auditoría de certificados SSL.
* **Kevin Retamales (Hardening Perimetral, Inteligencia pfBlockerNG y Control de Calidad QA):**
  - Hardening de políticas Zero Trust en pfSense y reglas de firewall flotantes.
  - Configuración de feeds de inteligencia de amenazas pfBlockerNG-devel (MaxMind GeoLite2 y listas FireHOL/Spamhaus).
  - Diseño y ejecución de matrices de pruebas funcionales de aseguramiento de calidad (QA).
  - Verificación de registros forenses EVE JSON y telemetría de red.

---

## 6. Evidencias

| Tipo de evidencia (avance o final) | Nombre de la evidencia | Descripción | Justificación |
| :--- | :--- | :--- | :--- |
| **Avance (Fase 1)** | Informe de Definición y Topología de Red | Documento técnico formal con diseño L2/L3, matriz de direccionamiento IP y fundamentación metodológica. | Valida la correcta planificación de la infraestructura de red perimetral y el marco teórico de partida. |
| **Avance (Fase 2)** | Repositorio Git de Código Fuente e IaC | Repositorio GitHub con el código del Motor KRONOS, Dockerfiles de Asterisk/DVWA y XML de pfSense. | Demuestra el avance tangible en desarrollo de software, integración de APIs y automatización de red. |
| **Avance (Fase 2)** | Manuales Técnicos y Mockups en PDF | Manual de configuración de pfSense CE 2.9.0 y Tutorial Paso a Paso con WebGUI mockups de alta calidad. | Acredita la reproducibilidad y rigor técnico de los procedimientos de hardening perimetral. |
| **Final (Fase 3)** | Demostración en Vivo de Intrusión y Respuesta SOAR | Ejecución de inyección SQL en DVWA ➔ Bloqueo en kernel FreeBSD (`snort2c`) ➔ Llamada por voz Gemini Live. | Evidencia concluyente del funcionamiento operativo y cumplimiento del objetivo general ante la comisión. |
| **Final (Fase 3)** | Matrices de Pruebas de Calidad (QA) y Logs Forenses | Registros EVE JSON de Suricata, salidas `pfctl -t snort2c -T show` y mediciones de latencia (<1.5 s). | Respaldan con datos empíricos la supresión de falsos positivos y la robustez del filtrado en kernel. |
| **Final (Fase 3)** | Informe Final Consolidado y Presentación Ejecutiva | Documento consolidado del proyecto y diapositivas de defensa ejecutiva en formato 16:9 widescreen. | Entrega formal académica exigida por la pauta de evaluación de Portafolio de Título (APT122). |

---

## 7. Plan de Trabajo

| Competencia o unidades de competencias | Nombre de Actividades/Tareas | Descripción Actividades/Tareas | Recursos | Duración de la actividad | Responsable¹ | Observaciones (Dificultades / Facilitadores) |
| :--- | :--- | :--- | :--- | :--- | :--- | :--- |
| **Comp. 4 & 8** | A1. Diseño y Setup Base pfSense | Instalación de pfSense CE 2.9.0, configuración WAN/LAN y tuning de hardware offloading para Netmap. | Hipervisor Proxmox/VMware, ISO pfSense CE 2.9.0 | Semanas 1 - 2 | Bruno Urrea / Freddy Vásquez | *Facilitador:* Documentación oficial Netgate.<br>*Dificultad:* Incompatibilidad de Netmap con TSO/LRO; se resuelve desactivando hardware offloading. |
| **Comp. 4** | A2. Segmentación de VLANs 802.1Q | Creación de subredes VLAN 10 (Corp), 20 (DMZ), 30 (VoIP) y 99 (Mgmt) y servidores DHCP locales. | pfSense WebGUI, Switch L2 virtual, perfiles VLAN | Semanas 3 - 4 | Freddy Vásquez | *Facilitador:* Soporte nativo 802.1Q en pfSense.<br>*Dificultad:* Filtrado inter-VLAN; se soluciona con reglas Zero Trust por interfaz. |
| **Comp. 7 & 8** | A3. Despliegue de Suricata Inline IPS | Instalación de Suricata 7.x, activación de modo Inline Netmap, reglas ET Open y configuración `dropsid.conf`. | Paquete Suricata, feeds ET Open, consola pfSense | Semanas 5 - 6 | Bruno Urrea / Kevin Retamales | *Facilitador:* Netmap permite descarte en ring-buffer sin latencia.<br>*Dificultad:* Ajuste de falsos positivos; se mitiga con reglas SID selectivas. |
| **Comp. 7 & 8** | A4. Hardening GeoIP con pfBlockerNG | Configuración de cuenta MaxMind Free, bloqueo Top Spammers y feeds FireHOL L1 / Spamhaus DROP. | pfBlockerNG-devel, cuenta MaxMind GeoLite2 | Semanas 7 - 8 | Kevin Retamales | *Facilitador:* Listas de reputación globales actualizadas.<br>*Dificultad:* Sobrecarga de memoria RAM; se optimiza el límite de tablas en pfSense. |
| **Comp. 7** | A5. Proxy Inverso HAProxy & DVWA | Configuración de Frontend HTTPS VIP 443, SSL Offloading, Stick-Tables anti-fuzzing y contenedor DVWA. | Paquete HAProxy, Docker Engine, imagen DVWA | Semanas 9 - 10 | Cristóbal Quezada | *Facilitador:* Stick-Tables en RAM procesan peticiones a nivel microsegundo.<br>*Dificultad:* Certificados autofirmados; se emite CA interna para el laboratorio. |
| **Comp. 6 & 8** | A6. Motor de Correlación KRONOS | Programación en Python 3.12 del parser `eve.json`, filtro heurístico AST y wrapper de kernel FreeBSD `pfctl`. | Python 3.12, librería PyYAML, FreeBSD CLI | Semanas 11 - 12 | Bruno Urrea | *Facilitador:* Librería estándar AST de Python.<br>*Dificultad:* Privilegios de ejecución en pfSense; se configura sudoers restringido para pfctl. |
| **Comp. 5** | A7. Centralita VoIP Asterisk PBX | Construcción de imagen Docker Asterisk 20 LTS, configuración de `pjsip.conf`, `extensions.conf` y auto-dialer AMI. | Docker Engine, Asterisk 20 LTS, AMI, softphone | Semanas 13 - 14 | Freddy Vásquez | *Facilitador:* Canal PJSIP moderno y ligero.<br>*Dificultad:* NAT traversal en VoIP; se resuelve con directivas `local_net` y `external_media_address`. |
| **Comp. 3 & 5** | A8. Integración Gemini Live Voice | Desarrollo de cliente WebSocket seguro, diseño de System Prompts tácticos y conexión de audio PCM con Asterisk. | Google AI Studio API Key, Python websockets | Semanas 14 - 15 | Bruno Urrea | *Facilitador:* Google Gemini Live Flash 3.1 Free Tier con baja latencia.<br>*Dificultad:* Sincronización dúplex de audio; se utiliza códec PCM lineal 24kHz. |
| **Comp. 3 & 4** | A9. Enlace Zero Trust Tailscale | Publicación de subred VoIP `192.168.30.0/24` en Tailscale para softphones móviles remotos sin abrir puertos. | Paquete Tailscale, túnel WireGuard Mesh | Semana 15 | Freddy Vásquez | *Facilitador:* WireGuard supera cualquier CGNAT o firewall intermedio.<br>*Dificultad:* Enrutamiento de subredes; se aprueba la ruta en el panel admin. |
| **Comp. 7, 8 & 11** | A10. Pruebas QA, Auditoría & Defensa | Ejecución de matrices de prueba de penetración, medición de tiempos (<1.5 s), manuales PDF y preparación de defensa. | Repositorio GitHub, softphone, ReportLab | Semanas 16 - 18 | Todo el Equipo | *Facilitador:* Roles bien delimitados y automatización previa.<br>*Dificultad:* Coordinación de demostración en vivo; se preparan scripts de contingencia. |

*¹ En caso de que el Proyecto APT sea grupal, en esta columna se indica el nombre de los responsables de cada tarea o actividad para permitir diferenciar la evaluación por cada integrante.*

---

## 8. Carta Gantt

| Actividad / Hito | Fase 1 (S1-S4) | | | | Fase 2 (S5-S15) | | | | | | | | | | | Fase 3 (S16-S18) | | |
| :--- | :---: | :---: | :---: | :---: | :---: | :---: | :---: | :---: | :---: | :---: | :---: | :---: | :---: | :---: | :---: | :---: | :---: | :---: | :---: |
| | **S1** | **S2** | **S3** | **S4** | **S5** | **S6** | **S7** | **S8** | **S9** | **S10** | **S11** | **S12** | **S13** | **S14** | **S15** | **S16** | **S17** | **S18** |
| **A1. Setup Base pfSense & Tuning** | **X** | **X** | | | | | | | | | | | | | | | | |
| **A2. Segmentación VLANs 802.1Q** | | | **X** | **X** | | | | | | | | | | | | | | |
| **A3. Suricata Inline Netmap IPS** | | | | | **X** | **X** | | | | | | | | | | | | |
| **A4. Hardening GeoIP pfBlockerNG** | | | | | | | **X** | **X** | | | | | | | | | | |
| **A5. HAProxy SSL & Laboratorio DVWA**| | | | | | | | | **X** | **X** | | | | | | | | |
| **A6. Motor de Correlación KRONOS**| | | | | | | | | | | **X** | **X** | | | | | | |
| **A7. Telefonía Asterisk PBX & AMI** | | | | | | | | | | | | | **X** | **X** | | | | |
| **A8. Integración Gemini Live Voice** | | | | | | | | | | | | | | **X** | **X** | | | |
| **A9. Malla Zero Trust Tailscale** | | | | | | | | | | | | | | | **X** | | | |
| **A10. Pruebas QA, Auditoría & Defensa**| | | | | | | | | | | | | | | | **X** | **X** | **X** |
"""

# =============================================================================
# 2. GENERADOR DE WORD (.DOCX) CON PORTADA, ÍNDICE Y FORMATO DUOC UC
# =============================================================================
def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def generate_docx(output_path: str):
    doc = docx.Document()
    
    # Configuración de márgenes estándar Duoc UC (2 cm)
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    # -------------------------------------------------------------
    # PORTADA INSTITUCIONAL FORMAL
    # -------------------------------------------------------------
    p_inst = doc.add_paragraph()
    p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_inst1 = p_inst.add_run("DUOC UC\nESCUELA DE INFORMÁTICA Y TELECOMUNICACIONES\n")
    r_inst1.font.name = "Arial"
    r_inst1.font.size = Pt(13)
    r_inst1.font.bold = True
    r_inst1.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    
    r_inst2 = p_inst.add_run("CARRERA DE INGENIERÍA EN CONECTIVIDAD Y REDES\nSEDE SAN JOAQUÍN\n")
    r_inst2.font.name = "Arial"
    r_inst2.font.size = Pt(11)
    r_inst2.font.bold = True
    r_inst2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph("\n" * 3)

    p_main_title = doc.add_paragraph()
    p_main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_mt = p_main_title.add_run("GUÍA 1: DEFINICIÓN PROYECTO APT\nASIGNATURA CAPSTONE (APT122) — FASE 1\n\n")
    r_mt.font.name = "Arial"
    r_mt.font.size = Pt(16)
    r_mt.font.bold = True
    r_mt.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    r_proj = p_main_title.add_run("PROYECTO:\nKRONOS SENTINEL: AUTONOMOUS AI-IPS & REAL-TIME INCIDENT VOICE RESPONSE SOAR ARCHITECTURE")
    r_proj.font.name = "Arial"
    r_proj.font.size = Pt(12)
    r_proj.font.bold = True
    r_proj.font.color.rgb = RGBColor(0x00, 0x66, 0x99)

    doc.add_paragraph("\n" * 4)

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    meta_runs = [
        ("Estudiante Titulación: ", True), ("Bruno Urrea Ortiz\n", False),
        ("RUT: ", True), ("21.543.637-3\n", False),
        ("Equipo de Trabajo: ", True), ("Bruno Urrea Ortiz\nFreddy Vásquez Cortés\nCristóbal Quezada\nKevin Retamales\n", False),
        ("Asignatura: ", True), ("Portafolio de Título (APT122)\n", False),
        ("Docente Evaluador: ", True), ("Comisión Evaluadora APT122\n", False),
        ("Fecha: ", True), ("Septiembre 2026\n", False)
    ]
    for text, bold in meta_runs:
        r = p_meta.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(10)
        r.font.bold = bold
        r.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    doc.add_page_break()

    # -------------------------------------------------------------
    # ÍNDICE DE CONTENIDOS Y TABLAS
    # -------------------------------------------------------------
    p_idx = doc.add_paragraph()
    r_idx = p_idx.add_run("ÍNDICE DE CONTENIDOS")
    r_idx.font.name = "Arial"
    r_idx.font.size = Pt(14)
    r_idx.font.bold = True
    r_idx.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    index_items = [
        ("1. Antecedentes Personales", "Pág. 3"),
        ("2. Descripción Proyecto APT (Áreas de Desempeño y Competencias)", "Pág. 3"),
        ("3. Fundamentación Proyecto APT", "Pág. 4"),
        ("   3.1 Relevancia del Proyecto APT", "Pág. 4"),
        ("   3.2 Descripción del Proyecto APT", "Pág. 4"),
        ("   3.3 Pertinencia del Proyecto con el Perfil de Egreso", "Pág. 4"),
        ("   3.4 Relación con los Intereses Profesionales", "Pág. 4"),
        ("   3.5 Factibilidad de Desarrollo del Proyecto APT", "Pág. 5"),
        ("4. Objetivos (Objetivo General y Específicos como Acciones)", "Pág. 5"),
        ("5. Metodología (Ciclo Iterativo y Roles del Equipo)", "Pág. 6"),
        ("6. Tabla Oficial de Evidencias (Avance y Finales)", "Pág. 7"),
        ("7. Tabla Oficial de Plan de Trabajo (Recursos y Observaciones)", "Pág. 8"),
        ("8. Carta Gantt (18 Semanas Académicas)", "Pág. 9")
    ]
    for item, page_str in index_items:
        p_item = doc.add_paragraph()
        r1 = p_item.add_run(item)
        r1.font.name = "Arial"
        r1.font.size = Pt(9.5)
        r1.font.bold = not item.startswith("   ")
        r2 = p_item.add_run(f" {'·' * (60 - len(item))} {page_str}")
        r2.font.name = "Arial"
        r2.font.size = Pt(9)
        r2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    # -------------------------------------------------------------
    # SECCIÓN 1: ANTECEDENTES PERSONALES (PARTE I)
    # -------------------------------------------------------------
    p_p1 = doc.add_paragraph()
    r_p1 = p_p1.add_run("A. PARTE I\n1. Antecedentes Personales")
    r_p1.font.name = "Arial"
    r_p1.font.size = Pt(12)
    r_p1.font.bold = True
    r_p1.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    
    t_pers = doc.add_table(rows=5, cols=2)
    t_pers.alignment = WD_TABLE_ALIGNMENT.CENTER
    pers_data = [
        ("Nombre estudiante", "Bruno Urrea Ortiz"),
        ("Rut", "21.543.637-3"),
        ("Carrera", "Ingeniería en Conectividad y Redes"),
        ("Sede", "San Joaquín"),
        ("Integrantes del Equipo (Grupal)", "• Bruno Urrea Ortiz (Líder de Ciberseguridad, Motor KRONOS y Gemini Live)\n• Freddy Vásquez Cortés (Routing, Switching y Telefonía Asterisk PBX)\n• Cristóbal Quezada (Servicios Web, HAProxy SSL y DMZ DVWA)\n• Kevin Retamales (Hardening Perimetral, pfBlockerNG y QA)")
    ]
    for idx, (k, v) in enumerate(pers_data):
        row = t_pers.rows[idx]
        row.cells[0].text = k
        row.cells[1].text = v
        set_cell_background(row.cells[0], "F0F4F8")
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        for cell in row.cells:
            set_cell_margins(cell)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(9.5)
                    
    doc.add_paragraph()

    # -------------------------------------------------------------
    # SECCIÓN 2: DESCRIPCIÓN PROYECTO APT
    # -------------------------------------------------------------
    p_s2 = doc.add_paragraph()
    r_s2 = p_s2.add_run("2. Descripción Proyecto APT")
    r_s2.font.name = "Arial"
    r_s2.font.size = Pt(12)
    r_s2.font.bold = True
    r_s2.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    
    t_desc = doc.add_table(rows=3, cols=2)
    t_desc.alignment = WD_TABLE_ALIGNMENT.CENTER
    desc_data = [
        ("Nombre del proyecto", "KRONOS SENTINEL: Autonomous AI-IPS & Real-Time Incident Voice Response SOAR Architecture"),
        ("Área(s) de desempeño", "• Seguridad de Sistemas y Redes: Ciberseguridad Defensiva y Hardening Perimetral.\n• Administración de Redes y Telecomunicaciones: Infraestructura L2/L3 (VLANs 802.1Q).\n• Comunicación Unificada y Corporativa: Telefonía VoIP SIP/PJSIP y Streaming de Voz.\n• Análisis de Soluciones de Conectividad: Automatización en Python 3.12 y SOAR con IA."),
        ("Competencias", "• Competencia 8: Crear planes de prevención y respuesta a riesgos informáticos en la red.\n• Competencia 6: Automatizar procesos y gestión de plataformas de red mediante scripting.\n• Competencia 5: Unificar servicios de voz, datos y video asegurando calidad de servicio (QoS).\n• Competencia 4: Controlar y operar redes corporativas de gran tamaño.\n• Competencia 3: Adaptar tecnologías de punta y tendencias tecnológicas emergentes.\n• Competencia 7: Gestionar la seguridad de la información frente a vulnerabilidades.")
    ]
    for idx, (k, v) in enumerate(desc_data):
        row = t_desc.rows[idx]
        row.cells[0].text = k
        row.cells[1].text = v
        set_cell_background(row.cells[0], "F0F4F8")
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        for cell in row.cells:
            set_cell_margins(cell)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(9.5)
                    
    doc.add_page_break()

    # -------------------------------------------------------------
    # SECCIÓN 3: FUNDAMENTACIÓN PROYECTO APT
    # -------------------------------------------------------------
    p_s3 = doc.add_paragraph()
    r_s3 = p_s3.add_run("3. Fundamentación Proyecto APT")
    r_s3.font.name = "Arial"
    r_s3.font.size = Pt(12)
    r_s3.font.bold = True
    r_s3.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    
    t_fund = doc.add_table(rows=5, cols=2)
    t_fund.alignment = WD_TABLE_ALIGNMENT.CENTER
    fund_data = [
        ("Relevancia del proyecto APT", "El proyecto resuelve dos fallas críticas en los SOC modernos:\n1) La crisis de más de 50% de falsos positivos en firewalls que satura las colas de incidentes.\n2) El colapso cognitivo del operador humano bajo ataque grave (visión de túnel entre aislar el firewall por CLI y contactar al CISO).\nKRONOS SENTINEL aporta una solución SOAR de Costo Cero ($0 CLP) que desacopla la contención atómica en kernel de FreeBSD (<100 ms) de la comunicación estratégica telefónica mediante Inteligencia Artificial conversacional (<1.5 s)."),
        ("Descripción del Proyecto APT", "Arquitectura de defensa en profundidad de 4 módulos:\n1) Firewall pfSense CE 2.9.0 con Suricata Inline Netmap IPS y pfBlockerNG GeoIP MaxMind.\n2) DMZ con HAProxy 2.8+ SSL Offloading y Stick-Tables anti-fuzzing L7 protegiendo DVWA en Docker.\n3) Motor de Correlación KRONOS (Python 3.12) con analizador heurístico AST y control de kernel con pfctl (terminación de estados y tabla snort2c).\n4) Centralita Asterisk 20 LTS en Docker con auto-dialer AMI y Google Gemini Live Flash 3.1 para debriefing interactivo por voz."),
        ("Pertinencia del proyecto con el perfil de egreso", "Integra armónicamente las 4 áreas troncales de Ingeniería en Conectividad y Redes: Routing & Switching L2/L3 (VLANs 802.1Q), Ciberseguridad Perimetral (Hardening, IPS, Zero Trust), Telefonía VoIP (Asterisk SIP/PJSIP) y Automatización (Python, IPC, IA Generativa)."),
        ("Relación con los intereses profesionales", "Se alinea directamente con mi meta profesional de desempeñarme como Ingeniero de Ciberseguridad / Arquitecto SecOps y SOAR, consolidando la experiencia práctica liderando el equipo de competiciones CTF DevSec y proyectos de mitigación perimetral con estándares internacionales (NIST SP 800-207, Cisco CCNA, CompTIA Security+)."),
        ("Factibilidad de desarrollo del Proyecto APT", "Alcance acotado y 100% factible para 18 semanas con 4 integrantes:\n• 18 semanas semestrales (72h taller presencial + 144h laboratorio autónomo = 864h de ingeniería grupal).\n• Materiales con Costo Cero ($0 CLP) basados en software libre (FreeBSD, pfSense, Suricata, Docker, Asterisk, Python) y capas gratuitas comunitarias (Gemini Live Free Tier, MaxMind GeoLite2).\n• Infraestructura de virtualización Proxmox VE / VMware y laboratorios Duoc UC.\n• Mitigación de riesgos: CGNAT resuelto mediante malla Zero Trust con Tailscale Subnet Router; falsos positivos controlados con AST; latencia de voz mitigada con streaming WebSocket PCM 24kHz.")
    ]
    for idx, (k, v) in enumerate(fund_data):
        row = t_fund.rows[idx]
        row.cells[0].text = k
        row.cells[1].text = v
        set_cell_background(row.cells[0], "F0F4F8")
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        for cell in row.cells:
            set_cell_margins(cell)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(9.2)

    doc.add_page_break()

    # -------------------------------------------------------------
    # SECCIÓN 4: OBJETIVOS (PARTE II)
    # -------------------------------------------------------------
    p_s4 = doc.add_paragraph()
    r_s4 = p_s4.add_run("B. PARTE II\n4. Objetivos")
    r_s4.font.name = "Arial"
    r_s4.font.size = Pt(12)
    r_s4.font.bold = True
    r_s4.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    
    t_obj = doc.add_table(rows=2, cols=2)
    t_obj.alignment = WD_TABLE_ALIGNMENT.CENTER
    obj_data = [
        ("Objetivo general", "Diseñar, implementar y validar una arquitectura de defensa en profundidad y respuesta autónoma ante incidentes (SOAR) de costo cero ($0 CLP) denominada KRONOS SENTINEL, integrando prevención de intrusiones en kernel, supresión heurística de falsos positivos e interlocución telefónica interactiva por voz en tiempo real con Inteligencia Artificial hacia los responsables de seguridad (CISO)."),
        ("Objetivos específicos", "1. Diseñar la topología lógica y física de red perimetral segmentada en 4 VLANs 802.1Q en pfSense CE 2.9.0.\n2. Implementar la configuración de pfSense CE 2.9.0 con políticas Zero Trust y optimización de hardware offloading para Netmap.\n3. Configurar el motor de prevención Suricata 7.x en modo Inline IPS (Netmap) con reglas dropsid.conf para descarte en hardware.\n4. Implementar el filtrado geográfico y reputacional con pfBlockerNG-devel utilizando MaxMind GeoLite2 y feeds FireHOL.\n5. Configurar el proxy inverso HAProxy 2.8+ SSL con Stick-Tables anti-fuzzing L7 protegiendo el entorno DVWA en DMZ.\n6. Desarrollar en Python 3.12 el Motor de Correlación KRONOS con análisis sintáctico AST suprimiendo >50% de falsos positivos.\n7. Integrar el control en kernel de FreeBSD mediante pfctl para terminación de estados (pfctl -k) y verificación en tabla snort2c.\n8. Desplegar la centralita Asterisk 20 LTS en Docker con canal PJSIP y módulo auto-dialer AMI hacia el softphone del CISO.\n9. Construir el puente WebSocket hacia Google Gemini Live Flash 3.1 para debriefing hablado bidireccional en tiempo real.\n10. Validar el rendimiento, latencia (<1.5 s) y efectividad integral mediante matrices de aseguramiento de calidad (QA).")
    ]
    for idx, (k, v) in enumerate(obj_data):
        row = t_obj.rows[idx]
        row.cells[0].text = k
        row.cells[1].text = v
        set_cell_background(row.cells[0], "F0F4F8")
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        for cell in row.cells:
            set_cell_margins(cell)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(9.2)

    doc.add_paragraph()

    # -------------------------------------------------------------
    # SECCIÓN 5: METODOLOGÍA
    # -------------------------------------------------------------
    p_s5 = doc.add_paragraph()
    r_s5 = p_s5.add_run("5. Metodología")
    r_s5.font.name = "Arial"
    r_s5.font.size = Pt(12)
    r_s5.font.bold = True
    r_s5.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    
    t_met = doc.add_table(rows=1, cols=1)
    t_met.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_met = t_met.rows[0].cells[0]
    cell_met.text = (
        "Descripción de la Metodología:\n"
        "Se aplicará una Metodología de Ingeniería en Ciclo Iterativo e Incremental en 4 fases operativas:\n"
        "• Fase 1: Levantamiento, Diseño y Topología de Red (Semanas 1 a 4).\n"
        "• Fase 2: Implementación de Perímetro y Capa Web DMZ (Semanas 5 a 10).\n"
        "• Fase 3: Desarrollo de KRONOS, Telefonía PBX y Voz IA (Semanas 11 a 15).\n"
        "• Fase 4: Pruebas de Integración SOAR, Auditoría QA y Defensa (Semanas 16 a 18).\n\n"
        "Distribución de Funciones, Tareas y Responsabilidades del Equipo:\n"
        "• Bruno Urrea Ortiz: Líder de Ciberseguridad, arquitectura global, desarrollo del Motor KRONOS AST, wrappers pfctl en kernel y WebSocket Gemini Live.\n"
        "• Freddy Vásquez Cortés: Ingeniería de Routing, segmentación VLAN 802.1Q, centralita Asterisk 20 LTS en Docker, auto-dialer AMI y malla Tailscale.\n"
        "• Cristóbal Quezada: Administración de Servicios Web, configuración HAProxy 2.8+ SSL, Stick-Tables de rate-limiting y laboratorio DVWA en DMZ.\n"
        "• Kevin Retamales: Hardening de firewall pfSense Zero Trust, pfBlockerNG-devel GeoIP y matrices de control de calidad QA."
    )
    set_cell_margins(cell_met)
    for p in cell_met.paragraphs:
        for r in p.runs:
            r.font.name = "Arial"
            r.font.size = Pt(9.2)

    doc.add_page_break()

    # -------------------------------------------------------------
    # SECCIÓN 6: EVIDENCIAS (TABLA OFICIAL 4 COLUMNAS)
    # -------------------------------------------------------------
    p_s6 = doc.add_paragraph()
    r_s6 = p_s6.add_run("6. Evidencias")
    r_s6.font.name = "Arial"
    r_s6.font.size = Pt(12)
    r_s6.font.bold = True
    r_s6.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    
    t_evi = doc.add_table(rows=7, cols=4)
    t_evi.alignment = WD_TABLE_ALIGNMENT.CENTER
    evi_headers = ["Tipo de evidencia (avance o final)", "Nombre de la evidencia", "Descripción", "Justificación"]
    for c_idx, h in enumerate(evi_headers):
        t_evi.rows[0].cells[c_idx].text = h
        set_cell_background(t_evi.rows[0].cells[c_idx], "003366")
        t_evi.rows[0].cells[c_idx].paragraphs[0].runs[0].font.bold = True
        t_evi.rows[0].cells[c_idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    evi_rows = [
        ("Avance (Fase 1)", "Informe de Definición y Topología de Red", "Documento técnico formal con diseño L2/L3, matriz de direccionamiento IP y fundamentación metodológica.", "Valida la correcta planificación de la infraestructura y el marco teórico."),
        ("Avance (Fase 2)", "Repositorio Git de Código Fuente e IaC", "Repositorio GitHub con el código del Motor KRONOS, Dockerfiles y configuraciones pfSense.", "Demuestra el avance tangible en desarrollo, integración de APIs y automatización."),
        ("Avance (Fase 2)", "Manuales Técnicos y Mockups en PDF", "Manual de configuración de pfSense CE 2.9.0 y Tutorial Paso a Paso con WebGUI mockups.", "Acredita la reproducibilidad y rigor técnico de los procedimientos perimetrales."),
        ("Final (Fase 3)", "Demostración en Vivo de Intrusión y SOAR", "Inyección SQL en DVWA ➔ Bloqueo en kernel FreeBSD (snort2c) ➔ Llamada por voz Gemini Live.", "Evidencia concluyente del funcionamiento operativo y cumplimiento del objetivo general."),
        ("Final (Fase 3)", "Matrices de Pruebas QA y Logs Forenses", "Registros EVE JSON de Suricata, salidas pfctl y mediciones de latencia (<1.5 s).", "Respaldan con datos empíricos la supresión de ruido y robustez en kernel."),
        ("Final (Fase 3)", "Informe Final Consolidado y Presentación", "Documento consolidado del proyecto y diapositivas de defensa ejecutiva en formato 16:9.", "Entrega formal académica exigida para la titulación en Portafolio APT122.")
    ]
    for r_idx, row_data in enumerate(evi_rows):
        row = t_evi.rows[r_idx+1]
        for c_idx, val in enumerate(row_data):
            row.cells[c_idx].text = val
            if r_idx % 2 == 1:
                set_cell_background(row.cells[c_idx], "F7FAFC")
            for p in row.cells[c_idx].paragraphs:
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(8.5)
            set_cell_margins(row.cells[c_idx], top=60, bottom=60, left=100, right=100)

    doc.add_page_break()

    # -------------------------------------------------------------
    # SECCIÓN 7: PLAN DE TRABAJO (TABLA OFICIAL 7 COLUMNAS)
    # -------------------------------------------------------------
    p_s7 = doc.add_paragraph()
    r_s7 = p_s7.add_run("7. Plan de Trabajo")
    r_s7.font.name = "Arial"
    r_s7.font.size = Pt(12)
    r_s7.font.bold = True
    r_s7.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    
    t_plan = doc.add_table(rows=11, cols=7)
    t_plan.alignment = WD_TABLE_ALIGNMENT.CENTER
    plan_headers = [
        "Competencia o unidades",
        "Nombre de Actividades/Tareas",
        "Descripción Actividades/Tareas",
        "Recursos",
        "Duración",
        "Responsable¹",
        "Observaciones (Facilitadores / Obstáculos)"
    ]
    for c_idx, h in enumerate(plan_headers):
        t_plan.rows[0].cells[c_idx].text = h
        set_cell_background(t_plan.rows[0].cells[c_idx], "003366")
        t_plan.rows[0].cells[c_idx].paragraphs[0].runs[0].font.bold = True
        t_plan.rows[0].cells[c_idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    plan_rows = [
        ("Comp. 4 & 8", "A1. Setup Base pfSense", "Instalación pfSense 2.9.0, WAN/LAN y tuning hardware offloading.", "pfSense ISO, Proxmox", "S1 - S2", "Bruno / Freddy", "Facilitador: Docs Netgate.\nObstáculo: Incompatibilidad TSO/LRO; se desactiva."),
        ("Comp. 4", "A2. Segmentación VLANs", "Creación de VLANs 10, 20, 30, 99 y servidores DHCP locales.", "pfSense WebGUI, Switch L2", "S3 - S4", "Freddy Vásquez", "Facilitador: Soporte 802.1Q.\nObstáculo: Fugas inter-VLAN; se aíslan con reglas."),
        ("Comp. 7 & 8", "A3. Suricata Inline IPS", "Suricata 7.x Inline Netmap, reglas ET Open y dropsid.conf.", "Suricata pkg, ET Open", "S5 - S6", "Bruno / Kevin", "Facilitador: Hardware drop Netmap.\nObstáculo: Falsos positivos; tuning de SIDs."),
        ("Comp. 7 & 8", "A4. GeoIP pfBlockerNG", "MaxMind Free, Top Spammers y feeds FireHOL L1/Spamhaus.", "pfBlockerNG, MaxMind", "S7 - S8", "Kevin Retamales", "Facilitador: Feeds globales.\nObstáculo: Memoria RAM; se optimiza límite tablas."),
        ("Comp. 7", "A5. HAProxy SSL & DVWA", "HAProxy SSL 443, Stick-Tables anti-fuzzing y DVWA Docker.", "HAProxy, Docker, DVWA", "S9 - S10", "Cristóbal Q.", "Facilitador: Tablas en RAM.\nObstáculo: Certificados SSL; se emite CA interna."),
        ("Comp. 6 & 8", "A6. Motor KRONOS AST", "Python 3.12 AST parser, eve.json ingesta y wrapper pfctl.", "Python 3.12, FreeBSD CLI", "S11 - S12", "Bruno Urrea", "Facilitador: AST de Python.\nObstáculo: Privilegios; se ajusta sudoers pfctl."),
        ("Comp. 5", "A7. Asterisk PBX & AMI", "Docker Asterisk 20 LTS, pjsip.conf y auto-dialer AMI.", "Docker, Asterisk 20 LTS", "S13 - S14", "Freddy Vásquez", "Facilitador: PJSIP moderno.\nObstáculo: NAT traversal; se fija external_media_address."),
        ("Comp. 3 & 5", "A8. Gemini Live Voice", "WebSocket seguro, System Prompts y audio bridge PCM 24kHz.", "Gemini Live API, Python", "S14 - S15", "Bruno Urrea", "Facilitador: Free tier baja latencia.\nObstáculo: Sincronía; códec PCM lineal."),
        ("Comp. 3 & 4", "A9. Malla Tailscale", "Publicación subred 192.168.30.0/24 en Tailscale Mesh.", "Tailscale pkg, WireGuard", "S15", "Freddy Vásquez", "Facilitador: WireGuard evade CGNAT.\nObstáculo: Rutas; aprobación manual panel."),
        ("Comp. 7, 8 & 11", "A10. Pruebas QA & Defensa", "Pruebas SQLi en vivo, manuales PDF y preparación defensa.", "ReportLab, softphone, Git", "S16 - S18", "Todo el Equipo", "Facilitador: Roles definidos.\nObstáculo: Contingencias; scripts de respaldo.")
    ]
    for r_idx, row_data in enumerate(plan_rows):
        row = t_plan.rows[r_idx+1]
        for c_idx, val in enumerate(row_data):
            row.cells[c_idx].text = val
            if r_idx % 2 == 1:
                set_cell_background(row.cells[c_idx], "F7FAFC")
            for p in row.cells[c_idx].paragraphs:
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(7.8)
            set_cell_margins(row.cells[c_idx], top=40, bottom=40, left=50, right=50)

    p_fn = doc.add_paragraph()
    r_fn = p_fn.add_run("¹ En caso de que el Proyecto APT sea grupal, en esta columna se indica el nombre de los responsables de cada tarea para diferenciar la evaluación.")
    r_fn.font.name = "Arial"
    r_fn.font.size = Pt(8)
    r_fn.font.italic = True

    doc.add_page_break()

    # -------------------------------------------------------------
    # SECCIÓN 8: CARTA GANTT (TABLA OFICIAL 18 SEMANAS)
    # -------------------------------------------------------------
    p_s8 = doc.add_paragraph()
    r_s8 = p_s8.add_run("8. Carta Gantt (18 Semanas Académicas)")
    r_s8.font.name = "Arial"
    r_s8.font.size = Pt(12)
    r_s8.font.bold = True
    r_s8.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    
    t_gantt = doc.add_table(rows=11, cols=19)
    t_gantt.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    t_gantt.rows[0].cells[0].text = "Actividad"
    set_cell_background(t_gantt.rows[0].cells[0], "003366")
    t_gantt.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
    t_gantt.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    for s in range(1, 19):
        cell = t_gantt.rows[0].cells[s]
        cell.text = f"S{s}"
        set_cell_background(cell, "003366")
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].runs[0].font.size = Pt(7)

    gantt_activities = [
        ("A1. Setup Base pfSense", [1, 2]),
        ("A2. Segmentación VLANs", [3, 4]),
        ("A3. Suricata Inline IPS", [5, 6]),
        ("A4. GeoIP pfBlockerNG", [7, 8]),
        ("A5. HAProxy SSL & DVWA", [9, 10]),
        ("A6. Motor KRONOS AST", [11, 12]),
        ("A7. Asterisk PBX & AMI", [13, 14]),
        ("A8. Gemini Live Voice", [14, 15]),
        ("A9. Malla Tailscale", [15]),
        ("A10. Pruebas QA & Defensa", [16, 17, 18])
    ]
    for r_idx, (act_name, weeks) in enumerate(gantt_activities):
        row = t_gantt.rows[r_idx+1]
        row.cells[0].text = act_name
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(7.5)
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        set_cell_margins(row.cells[0], top=30, bottom=30, left=50, right=50)
        
        for w in range(1, 19):
            cell = row.cells[w]
            set_cell_margins(cell, top=30, bottom=30, left=20, right=20)
            if w in weeks:
                cell.text = "X"
                set_cell_background(cell, "006699")
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.size = Pt(7)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                cell.text = ""

    doc.save(output_path)
    print(f"[OK] Word DOCX generated with Portada & Índice: {output_path}")

# =============================================================================
# 3. GENERADOR DE PDF OFICIAL CON REPORTLAB (PORTADA, ÍNDICE & TABLAS)
# =============================================================================
class DuocOfficialCanvas(canvas.Canvas):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            if self._pageNumber > 1:
                self.draw_header_footer(num_pages)
            super().showPage()
        super().save()

    def draw_header_footer(self, page_count):
        self.saveState()
        # Header institucional
        self.setFont("Helvetica-Bold", 8)
        self.setFillColor(colors.HexColor("#003366"))
        self.drawString(40, 755, "DUOC UC — ESCUELA DE INFORMÁTICA Y TELECOMUNICACIONES")
        
        self.setFont("Helvetica", 8)
        self.setFillColor(colors.HexColor("#666666"))
        self.drawRightString(572, 755, "Guía 1. Definición Proyecto APT (APT122)")
        
        self.setStrokeColor(colors.HexColor("#CCCCCC"))
        self.setLineWidth(0.5)
        self.line(40, 748, 572, 748)
        
        # Footer
        self.setFont("Helvetica", 8)
        self.setFillColor(colors.HexColor("#666666"))
        self.drawString(40, 30, "KRONOS SENTINEL // Estudiante: Bruno Urrea Ortiz (RUT: 21.543.637-3)")
        
        self.setFont("Helvetica-Bold", 8)
        self.setFillColor(colors.HexColor("#003366"))
        self.drawRightString(572, 30, f"Página {self._pageNumber} de {page_count}")
        
        self.setStrokeColor(colors.HexColor("#CCCCCC"))
        self.setLineWidth(0.5)
        self.line(40, 40, 572, 40)
        self.restoreState()

def generate_pdf(output_path: str):
    doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        leftMargin=40,
        rightMargin=40,
        topMargin=55,
        bottomMargin=55
    )

    styles = getSampleStyleSheet()
    
    st_cover_inst = ParagraphStyle('CInst', fontName='Helvetica-Bold', fontSize=14, leading=18, textColor=colors.HexColor("#003366"), alignment=1)
    st_cover_sub = ParagraphStyle('CSub', fontName='Helvetica-Bold', fontSize=11, leading=15, textColor=colors.HexColor("#555555"), alignment=1)
    st_cover_title = ParagraphStyle('CTitle', fontName='Helvetica-Bold', fontSize=18, leading=22, textColor=colors.HexColor("#003366"), alignment=1)
    st_cover_proj = ParagraphStyle('CProj', fontName='Helvetica-Bold', fontSize=12, leading=16, textColor=colors.HexColor("#006699"), alignment=1)
    st_cover_meta = ParagraphStyle('CMeta', fontName='Helvetica', fontSize=9.5, leading=14, textColor=colors.HexColor("#222222"), alignment=2)
    
    st_sec_title = ParagraphStyle('SecT', fontName='Helvetica-Bold', fontSize=12, leading=15, textColor=colors.HexColor("#003366"), spaceBefore=8, spaceAfter=4)
    st_sub_title = ParagraphStyle('SubT', fontName='Helvetica-Bold', fontSize=10, leading=13, textColor=colors.HexColor("#006699"), spaceBefore=6, spaceAfter=3)
    
    st_th = ParagraphStyle('TH', fontName='Helvetica-Bold', fontSize=8, leading=10, textColor=colors.white, alignment=1)
    st_td_label = ParagraphStyle('TDL', fontName='Helvetica-Bold', fontSize=8, leading=10.5, textColor=colors.HexColor("#003366"))
    st_td_val = ParagraphStyle('TDV', fontName='Helvetica', fontSize=8, leading=10.5, textColor=colors.HexColor("#222222"))
    st_td_val_mono = ParagraphStyle('TDVM', fontName='Helvetica-Bold', fontSize=7.5, leading=9.5, textColor=colors.HexColor("#003366"), alignment=1)

    story = []

    # -------------------------------------------------------------
    # PORTADA PDF
    # -------------------------------------------------------------
    story.append(Spacer(1, 20))
    story.append(Paragraph("DUOC UC<br/>ESCUELA DE INFORMÁTICA Y TELECOMUNICACIONES", st_cover_inst))
    story.append(Spacer(1, 4))
    story.append(Paragraph("CARRERA DE INGENIERÍA EN CONECTIVIDAD Y REDES<br/>SEDE SAN JOAQUÍN", st_cover_sub))
    story.append(Spacer(1, 60))
    story.append(Paragraph("GUÍA 1: DEFINICIÓN PROYECTO APT<br/>ASIGNATURA CAPSTONE (APT122) — FASE 1", st_cover_title))
    story.append(Spacer(1, 10))
    story.append(Paragraph("PROYECTO:<br/>KRONOS SENTINEL: AUTONOMOUS AI-IPS & REAL-TIME INCIDENT VOICE RESPONSE SOAR ARCHITECTURE", st_cover_proj))
    story.append(Spacer(1, 100))
    
    meta_text = """
    <b>Estudiante Titulación:</b> Bruno Urrea Ortiz<br/>
    <b>RUT:</b> 21.543.637-3<br/>
    <b>Equipo de Trabajo:</b> Bruno Urrea Ortiz, Freddy Vásquez Cortés,<br/>
    Cristóbal Quezada, Kevin Retamales<br/>
    <b>Asignatura:</b> Portafolio de Título (APT122)<br/>
    <b>Docente Evaluador:</b> Comisión Evaluadora APT122<br/>
    <b>Fecha:</b> Septiembre 2026
    """
    story.append(Paragraph(meta_text, st_cover_meta))
    story.append(PageBreak())

    # -------------------------------------------------------------
    # ÍNDICE PDF
    # -------------------------------------------------------------
    story.append(Paragraph("ÍNDICE DE CONTENIDOS", st_sec_title))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#003366"), spaceAfter=10))
    
    index_p_text = """
    <b>1. Antecedentes Personales</b> ............................................................................................ Pág. 3<br/>
    <b>2. Descripción Proyecto APT (Áreas de Desempeño y Competencias)</b> .......................... Pág. 3<br/>
    <b>3. Fundamentación Proyecto APT</b> ................................................................................. Pág. 3<br/>
    &nbsp;&nbsp;&nbsp;&nbsp;3.1 Relevancia del Proyecto APT .............................................................................. Pág. 3<br/>
    &nbsp;&nbsp;&nbsp;&nbsp;3.2 Descripción del Proyecto APT ........................................................................... Pág. 4<br/>
    &nbsp;&nbsp;&nbsp;&nbsp;3.3 Pertinencia del Proyecto con el Perfil de Egreso ................................................... Pág. 4<br/>
    &nbsp;&nbsp;&nbsp;&nbsp;3.4 Relación con los Intereses Profesionales ............................................................. Pág. 4<br/>
    &nbsp;&nbsp;&nbsp;&nbsp;3.5 Factibilidad de Desarrollo del Proyecto APT ...................................................... Pág. 4<br/>
    <b>4. Objetivos (General y Específicos como Acciones)</b> ...................................................... Pág. 5<br/>
    <b>5. Metodología (Ciclo Iterativo y Roles del Equipo)</b> ...................................................... Pág. 5<br/>
    <b>6. Tabla Oficial de Evidencias (Avance y Finales)</b> ......................................................... Pág. 6<br/>
    <b>7. Tabla Oficial de Plan de Trabajo (Recursos y Observaciones)</b> ................................. Pág. 6<br/>
    <b>8. Carta Gantt (18 Semanas Académicas)</b> ....................................................................... Pág. 7
    """
    story.append(Paragraph(index_p_text, st_td_val))
    story.append(PageBreak())

    # -------------------------------------------------------------
    # PARTE I: ANTECEDENTES, DESCRIPCIÓN & FUNDAMENTACIÓN
    # -------------------------------------------------------------
    story.append(Paragraph("A. PARTE I", st_sec_title))
    story.append(Paragraph("1. Antecedentes Personales", st_sub_title))
    
    t1_data = [
        [Paragraph("Nombre estudiante", st_td_label), Paragraph("Bruno Urrea Ortiz", st_td_val)],
        [Paragraph("Rut", st_td_label), Paragraph("21.543.637-3", st_td_val)],
        [Paragraph("Carrera", st_td_label), Paragraph("Ingeniería en Conectividad y Redes", st_td_val)],
        [Paragraph("Sede", st_td_label), Paragraph("San Joaquín", st_td_val)],
        [Paragraph("Integrantes del Equipo (Grupal)", st_td_label), Paragraph("• <b>Bruno Urrea Ortiz:</b> Líder Ciberseguridad, Motor KRONOS y Gemini Live.<br/>• <b>Freddy Vásquez Cortés:</b> Routing, Switching y Telefonía Asterisk PBX.<br/>• <b>Cristóbal Quezada:</b> Servicios Web, Proxy Inverso HAProxy SSL y DMZ.<br/>• <b>Kevin Retamales:</b> Hardening Perimetral, pfBlockerNG y QA.", st_td_val)]
    ]
    t1 = Table(t1_data, colWidths=[150, 382])
    t1.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (0,-1), colors.HexColor("#F0F4F8")),
        ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor("#CCCCCC")),
        ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor("#E2E8F0")),
        ('TOPPADDING', (0,0), (-1,-1), 3),
        ('BOTTOMPADDING', (0,0), (-1,-1), 3),
    ]))
    story.append(t1)
    story.append(Spacer(1, 8))

    story.append(Paragraph("2. Descripción Proyecto APT", st_sub_title))
    t2_data = [
        [Paragraph("Nombre del proyecto", st_td_label), Paragraph("<b>KRONOS SENTINEL:</b> Autonomous AI-IPS & Real-Time Incident Voice Response SOAR Architecture", st_td_val)],
        [Paragraph("Área(s) de desempeño", st_td_label), Paragraph("• Seguridad de Sistemas y Redes / Ciberseguridad Defensiva.<br/>• Administración de Redes y Telecomunicaciones / Infraestructura L2/L3.<br/>• Comunicación Unificada y Corporativa / Telefonía VoIP SIP/PJSIP.<br/>• Análisis de Soluciones de Conectividad / Python y SOAR con IA.", st_td_val)],
        [Paragraph("Competencias", st_td_label), Paragraph("• <b>Comp. 8:</b> Crear planes de prevención y respuesta a riesgos informáticos.<br/>• <b>Comp. 6:</b> Automatizar procesos y plataformas de red mediante scripting.<br/>• <b>Comp. 5:</b> Unificar servicios de voz, datos y video con calidad de servicio (QoS).<br/>• <b>Comp. 4:</b> Controlar y operar redes corporativas de gran tamaño.<br/>• <b>Comp. 3:</b> Adaptar tecnologías de punta y tendencias emergentes.<br/>• <b>Comp. 7:</b> Gestionar la seguridad frente a vulnerabilidades en aplicaciones.", st_td_val)]
    ]
    t2 = Table(t2_data, colWidths=[150, 382])
    t2.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (0,-1), colors.HexColor("#F0F4F8")),
        ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor("#CCCCCC")),
        ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor("#E2E8F0")),
        ('TOPPADDING', (0,0), (-1,-1), 3),
        ('BOTTOMPADDING', (0,0), (-1,-1), 3),
    ]))
    story.append(t2)
    story.append(Spacer(1, 8))

    story.append(Paragraph("3. Fundamentación Proyecto APT", st_sub_title))
    t3_data = [
        [Paragraph("Relevancia del proyecto APT", st_td_label), Paragraph("Resuelve dos fallas críticas en SOCs:<br/>1) <b>>50% falsos positivos</b> en IPS tradicionales.<br/>2) <b>Colapso cognitivo del operador</b> bajo ataque grave (visión de túnel entre consola de firewall y redactar reportes al CISO).<br/><b>Aporte:</b> SOAR $0 CLP que contiene en kernel FreeBSD (<100 ms) y llama con IA por voz (<1.5 s).", st_td_val)],
        [Paragraph("Descripción del Proyecto APT", st_td_label), Paragraph("4 pilares tecnológicos:<br/>1) pfSense 2.9.0 + Suricata Inline Netmap IPS + pfBlockerNG GeoIP.<br/>2) DMZ HAProxy SSL + Stick-Tables anti-fuzzing + DVWA en Docker.<br/>3) Motor KRONOS Python 3.12 (parser AST anti-ruido + pfctl kill states).<br/>4) Asterisk 20 LTS Docker + AMI auto-dialer + Gemini Live Flash 3.1.", st_td_val)],
        [Paragraph("Pertinencia con perfil de egreso", st_td_label), Paragraph("Integra armónicamente Networking L2/L3 (VLANs 802.1Q), Ciberseguridad Perimetral (Hardening FreeBSD, Netmap), Telefonía VoIP (Asterisk SIP) y Automatización (Python, IA Generativa).", st_td_val)],
        [Paragraph("Relación con intereses profesionales", st_td_label), Paragraph("Especialización en Ciberseguridad Defensiva, SecOps y SOAR. Consolida el liderazgo en CTFs con el equipo DevSec y alinea con certificaciones internacionales (Cisco CCNA, CompTIA Security+, NIST SP 800-207).", st_td_val)],
        [Paragraph("Factibilidad de desarrollo (18 Semanas)", st_td_label), Paragraph("<b>18 semanas / 4 integrantes (864h de ingeniería grupal):</b><br/>• $0 CLP: 100% Open Source y capas gratuitas (Gemini Live Free Tier, MaxMind).<br/>• Proxmox VE y labs Duoc UC.<br/>• Mitigaciones: Tailscale supera CGNAT; AST suprime falsos positivos; WebSocket PCM 24kHz reduce latencia.", st_td_val)]
    ]
    t3 = Table(t3_data, colWidths=[150, 382])
    t3.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (0,-1), colors.HexColor("#F0F4F8")),
        ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor("#CCCCCC")),
        ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor("#E2E8F0")),
        ('TOPPADDING', (0,0), (-1,-1), 3),
        ('BOTTOMPADDING', (0,0), (-1,-1), 3),
    ]))
    story.append(t3)
    story.append(PageBreak())

    # -------------------------------------------------------------
    # PARTE II: OBJETIVOS & METODOLOGÍA
    # -------------------------------------------------------------
    story.append(Paragraph("B. PARTE II", st_sec_title))
    story.append(Paragraph("4. Objetivos", st_sub_title))
    
    t4_data = [
        [Paragraph("Objetivo general", st_td_label), Paragraph("Diseñar, implementar y validar una arquitectura de defensa en profundidad y respuesta autónoma ante incidentes (SOAR) de costo cero ($0 CLP) denominada <b>KRONOS SENTINEL</b>, integrando prevención de intrusiones en kernel, supresión heurística de falsos positivos e interlocución telefónica interactiva por voz en tiempo real con Inteligencia Artificial hacia los responsables de seguridad (CISO).", st_td_val)],
        [Paragraph("Objetivos específicos (Acciones)", st_td_label), Paragraph(
            "1. <b>Diseñar</b> la topología lógica/física de red en 4 VLANs 802.1Q en pfSense 2.9.0.<br/>"
            "2. <b>Implementar</b> la configuración de pfSense con políticas Zero Trust y tuning Netmap.<br/>"
            "3. <b>Configurar</b> Suricata 7.x Inline IPS con reglas dropsid.conf para hardware drop.<br/>"
            "4. <b>Implementar</b> el filtrado GeoIP y listas FireHOL con pfBlockerNG-devel.<br/>"
            "5. <b>Configurar</b> HAProxy 2.8+ SSL y Stick-Tables protegiendo DVWA en DMZ.<br/>"
            "6. <b>Desarrollar</b> en Python 3.12 el Motor KRONOS con parser AST anti-ruido (>50%).<br/>"
            "7. <b>Integrar</b> el control en kernel FreeBSD con pfctl (kill states y tabla snort2c).<br/>"
            "8. <b>Desplegar</b> Asterisk 20 LTS en Docker con canal PJSIP y auto-dialer AMI.<br/>"
            "9. <b>Construir</b> el puente WebSocket hacia Gemini Live Flash 3.1 (PCM 24kHz).<br/>"
            "10. <b>Validar</b> el rendimiento, latencia (<1.5 s) y efectividad con matrices QA.",
            st_td_val
        )]
    ]
    t4 = Table(t4_data, colWidths=[150, 382])
    t4.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (0,-1), colors.HexColor("#F0F4F8")),
        ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor("#CCCCCC")),
        ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor("#E2E8F0")),
        ('TOPPADDING', (0,0), (-1,-1), 3),
        ('BOTTOMPADDING', (0,0), (-1,-1), 3),
    ]))
    story.append(t4)
    story.append(Spacer(1, 8))

    story.append(Paragraph("5. Metodología", st_sub_title))
    t5_data = [
        [Paragraph(
            "<b>Descripción de la Metodología (Ciclo Iterativo e Incremental en 4 Fases):</b><br/>"
            "• <b>Fase 1:</b> Levantamiento, Diseño y Topología de Red (Semanas 1 a 4).<br/>"
            "• <b>Fase 2:</b> Implementación de Perímetro y Capa Web DMZ (Semanas 5 a 10).<br/>"
            "• <b>Fase 3:</b> Desarrollo de KRONOS, Telefonía PBX y Voz IA (Semanas 11 a 15).<br/>"
            "• <b>Fase 4:</b> Pruebas de Integración SOAR, Auditoría QA y Defensa (Semanas 16 a 18).<br/><br/>"
            "<b>Distribución de Funciones y Responsabilidades:</b><br/>"
            "• <b>Bruno Urrea Ortiz:</b> Líder Ciberseguridad, Motor KRONOS AST, pfctl wrappers y Gemini Live.<br/>"
            "• <b>Freddy Vásquez Cortés:</b> Ingeniería de Routing, VLANs 802.1Q, Asterisk Docker, AMI y Tailscale.<br/>"
            "• <b>Cristóbal Quezada:</b> Servicios Web, HAProxy 2.8+ SSL, Stick-Tables y laboratorio DVWA.<br/>"
            "• <b>Kevin Retamales:</b> Hardening pfSense Zero Trust, pfBlockerNG GeoIP y matrices de QA.",
            st_td_val
        )]
    ]
    t5 = Table(t5_data, colWidths=[532])
    t5.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor("#F0F4F8")),
        ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor("#CCCCCC")),
        ('TOPPADDING', (0,0), (-1,-1), 5),
        ('BOTTOMPADDING', (0,0), (-1,-1), 5),
        ('LEFTPADDING', (0,0), (-1,-1), 8),
        ('RIGHTPADDING', (0,0), (-1,-1), 8),
    ]))
    story.append(t5)
    story.append(PageBreak())

    # -------------------------------------------------------------
    # SECCIÓN 6: EVIDENCIAS & SECCIÓN 7: PLAN DE TRABAJO
    # -------------------------------------------------------------
    story.append(Paragraph("6. Evidencias (Tabla Oficial)", st_sec_title))
    t6_data = [
        [Paragraph("Tipo de evidencia", st_th), Paragraph("Nombre de la evidencia", st_th), Paragraph("Descripción", st_th), Paragraph("Justificación", st_th)],
        [Paragraph("Avance (F1)", st_td_label), Paragraph("Informe Definición y Topología", st_td_val), Paragraph("Diseño L2/L3, VLANs y fundamentación.", st_td_val), Paragraph("Valida infraestructura inicial.", st_td_val)],
        [Paragraph("Avance (F2)", st_td_label), Paragraph("Repositorio Git e IaC", st_td_val), Paragraph("Código KRONOS, Dockerfiles, XML pfSense.", st_td_val), Paragraph("Acredita avance de desarrollo.", st_td_val)],
        [Paragraph("Avance (F2)", st_td_label), Paragraph("Manuales Técnicos PDF", st_td_val), Paragraph("Manual pfSense y Tutorial WebGUI.", st_td_val), Paragraph("Acredita reproducibilidad técnica.", st_td_val)],
        [Paragraph("Final (F3)", st_td_label), Paragraph("Demostración en Vivo SOAR", st_td_val), Paragraph("SQLi ➔ Drop Netmap ➔ pfctl ➔ Gemini Live.", st_td_val), Paragraph("Evidencia concluyente de la solución.", st_td_val)],
        [Paragraph("Final (F3)", st_td_label), Paragraph("Matrices QA y Logs Forenses", st_td_val), Paragraph("EVE JSON, tablas snort2c, latencia <1.5s.", st_td_val), Paragraph("Datos empíricos de mitigación.", st_td_val)],
        [Paragraph("Final (F3)", st_td_label), Paragraph("Informe Final y Presentación", st_td_val), Paragraph("Documento consolidado y slides 16:9.", st_td_val), Paragraph("Entrega formal Portafolio APT122.", st_td_val)]
    ]
    t6 = Table(t6_data, colWidths=[65, 145, 175, 147])
    t6.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#003366")),
        ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor("#CCCCCC")),
        ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor("#E2E8F0")),
        ('TOPPADDING', (0,0), (-1,-1), 3),
        ('BOTTOMPADDING', (0,0), (-1,-1), 3),
    ]))
    story.append(t6)
    story.append(Spacer(1, 10))

    story.append(Paragraph("7. Plan de Trabajo (Tabla Oficial)", st_sec_title))
    t7_data = [
        [Paragraph("Competencia", st_th), Paragraph("Actividad", st_th), Paragraph("Descripción", st_th), Paragraph("Recursos", st_th), Paragraph("Duración", st_th), Paragraph("Responsable¹", st_th), Paragraph("Observaciones", st_th)],
        [Paragraph("Comp. 4 & 8", st_td_val), Paragraph("A1. Setup pfSense", st_td_val), Paragraph("Setup pfSense 2.9.0 y tuning Netmap.", st_td_val), Paragraph("Proxmox, ISO", st_td_val), Paragraph("S1-S2", st_td_val), Paragraph("Bruno/Freddy", st_td_val), Paragraph("Desactivar TSO/LRO.", st_td_val)],
        [Paragraph("Comp. 4", st_td_val), Paragraph("A2. VLANs 802.1Q", st_td_val), Paragraph("VLANs 10, 20, 30, 99 y DHCP.", st_td_val), Paragraph("pfSense WebGUI", st_td_val), Paragraph("S3-S4", st_td_val), Paragraph("Freddy Vásquez", st_td_val), Paragraph("Reglas Zero Trust.", st_td_val)],
        [Paragraph("Comp. 7 & 8", st_td_val), Paragraph("A3. Suricata IPS", st_td_val), Paragraph("Suricata Netmap y dropsid.conf.", st_td_val), Paragraph("ET Open, pkg", st_td_val), Paragraph("S5-S6", st_td_val), Paragraph("Bruno/Kevin", st_td_val), Paragraph("Hardware drop ring-buffer.", st_td_val)],
        [Paragraph("Comp. 7 & 8", st_td_val), Paragraph("A4. GeoIP pfBlocker", st_td_val), Paragraph("MaxMind Free y listas FireHOL.", st_td_val), Paragraph("pfBlockerNG", st_td_val), Paragraph("S7-S8", st_td_val), Paragraph("Kevin Retamales", st_td_val), Paragraph("Optimizar RAM tablas.", st_td_val)],
        [Paragraph("Comp. 7", st_td_val), Paragraph("A5. HAProxy & DVWA", st_td_val), Paragraph("HAProxy SSL y DVWA Docker.", st_td_val), Paragraph("Docker, HAProxy", st_td_val), Paragraph("S9-S10", st_td_val), Paragraph("Cristóbal Q.", st_td_val), Paragraph("Stick-Tables en RAM.", st_td_val)],
        [Paragraph("Comp. 6 & 8", st_td_val), Paragraph("A6. Motor KRONOS", st_td_val), Paragraph("Python AST parser y pfctl wrapper.", st_td_val), Paragraph("Python 3.12", st_td_val), Paragraph("S11-S12", st_td_val), Paragraph("Bruno Urrea", st_td_val), Paragraph("Sudoers pfctl restringido.", st_td_val)],
        [Paragraph("Comp. 5", st_td_val), Paragraph("A7. Asterisk PBX", st_td_val), Paragraph("Asterisk Docker y auto-dialer AMI.", st_td_val), Paragraph("Asterisk 20", st_td_val), Paragraph("S13-S14", st_td_val), Paragraph("Freddy Vásquez", st_td_val), Paragraph("external_media_address.", st_td_val)],
        [Paragraph("Comp. 3 & 5", st_td_val), Paragraph("A8. Gemini Live", st_td_val), Paragraph("WebSocket audio bridge PCM 24kHz.", st_td_val), Paragraph("Gemini API", st_td_val), Paragraph("S14-S15", st_td_val), Paragraph("Bruno Urrea", st_td_val), Paragraph("Free Tier baja latencia.", st_td_val)],
        [Paragraph("Comp. 3 & 4", st_td_val), Paragraph("A9. Malla Tailscale", st_td_val), Paragraph("Subred 192.168.30.0/24 WireGuard.", st_td_val), Paragraph("Tailscale pkg", st_td_val), Paragraph("S15", st_td_val), Paragraph("Freddy Vásquez", st_td_val), Paragraph("Supera bloqueo CGNAT.", st_td_val)],
        [Paragraph("Comp. 7, 8, 11", st_td_val), Paragraph("A10. QA & Defensa", st_td_val), Paragraph("Pruebas QA <1.5s y manuales PDF.", st_td_val), Paragraph("ReportLab, Git", st_td_val), Paragraph("S16-S18", st_td_val), Paragraph("Todo el Equipo", st_td_val), Paragraph("Scripts contingencia.", st_td_val)]
    ]
    t7 = Table(t7_data, colWidths=[55, 75, 125, 65, 38, 65, 109])
    t7.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#003366")),
        ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor("#CCCCCC")),
        ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor("#E2E8F0")),
        ('TOPPADDING', (0,0), (-1,-1), 2),
        ('BOTTOMPADDING', (0,0), (-1,-1), 2),
    ]))
    story.append(t7)
    story.append(Spacer(1, 3))
    story.append(Paragraph("<i>¹ En caso de que el Proyecto APT sea grupal, en esta columna se indica el responsable de cada tarea para diferenciar la evaluación individual.</i>", st_td_val))
    story.append(PageBreak())

    # -------------------------------------------------------------
    # SECCIÓN 8: CARTA GANTT
    # -------------------------------------------------------------
    story.append(Paragraph("8. Carta Gantt (18 Semanas Académicas)", st_sec_title))
    
    gantt_hdr = [Paragraph("Actividad", st_th)] + [Paragraph(f"S{s}", st_th) for s in range(1, 19)]
    t8_data = [gantt_hdr]
    
    gantt_activities = [
        ("A1. Setup Base pfSense", [1, 2]),
        ("A2. Segmentación VLANs", [3, 4]),
        ("A3. Suricata Inline IPS", [5, 6]),
        ("A4. GeoIP pfBlockerNG", [7, 8]),
        ("A5. HAProxy SSL & DVWA", [9, 10]),
        ("A6. Motor KRONOS AST", [11, 12]),
        ("A7. Asterisk PBX & AMI", [13, 14]),
        ("A8. Gemini Live Voice", [14, 15]),
        ("A9. Malla Tailscale", [15]),
        ("A10. Pruebas QA & Defensa", [16, 17, 18])
    ]
    for act_name, weeks in gantt_activities:
        row = [Paragraph(act_name, st_td_val)]
        for s in range(1, 19):
            if s in weeks:
                row.append(Paragraph("X", st_td_val_mono))
            else:
                row.append(Paragraph("", st_td_val))
        t8_data.append(row)

    t8 = Table(t8_data, colWidths=[136] + [22]*18)
    t8_style = [
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#003366")),
        ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor("#CCCCCC")),
        ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor("#E2E8F0")),
        ('TOPPADDING', (0,0), (-1,-1), 2),
        ('BOTTOMPADDING', (0,0), (-1,-1), 2),
        ('ALIGN', (1,1), (-1,-1), 'CENTER'),
    ]
    # Highlight Gantt cells
    for r_idx, (act_name, weeks) in enumerate(gantt_activities):
        for w in weeks:
            t8_style.append(('BACKGROUND', (w, r_idx+1), (w, r_idx+1), colors.HexColor("#D0E8F2")))
            
    t8.setStyle(TableStyle(t8_style))
    story.append(t8)

    doc.build(story, canvasmaker=DuocOfficialCanvas)
    print(f"[OK] Official PDF generated with Portada & Índice: {output_path}")

# =============================================================================
# 4. EJECUTOR PRINCIPAL
# =============================================================================
def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    
    # 1. Guardar Markdown
    with open(MD_FILE, "w", encoding="utf-8") as f:
        f.write(MARKDOWN_CONTENT.strip() + "\n")
    print(f"[OK] Markdown generated: {MD_FILE}")

    # 2. Generar Word (.docx)
    generate_docx(DOCX_FILE)

    # 3. Generar PDF (.pdf)
    generate_pdf(PDF_FILE)

if __name__ == "__main__":
    main()
