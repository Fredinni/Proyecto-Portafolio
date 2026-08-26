import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def create_autoevaluacion_docx(filename):
    doc = docx.Document()
    
    # Title
    title = doc.add_paragraph()
    r = title.add_run("Pauta de Autoevaluación de Competencias\n")
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0, 51, 102)
    
    sub = title.add_run("(complemento de la Pauta de Reflexión Definición Proyecto APT)\nFase 1 - APT122")
    sub.font.size = Pt(11)
    sub.font.italic = True
    
    doc.add_paragraph("Objetivo: Identificar los niveles de logro en las competencias del plan de estudio de Ingeniería en Conectividad y Redes para definir y fundamentar el proyecto APT (Portafolio de Título).")

    # Table Info Estudiante
    t_info = doc.add_table(rows=4, cols=2)
    t_info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ("Escuela", "Escuela de Informática y Telecomunicaciones - Duoc UC (Sede San Joaquín)"),
        ("Nombre completo", "Bruno Urrea Ortiz"),
        ("Plan de Estudio", "Ingeniería en Conectividad y Redes"),
        ("Año de ingreso", "2022 / 2023")
    ]
    for i, (k, v) in enumerate(info_data):
        row = t_info.rows[i]
        row.cells[0].text = k
        row.cells[1].text = v
        set_cell_background(row.cells[0], "F0F4F8")
        row.cells[0].paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph()

    # Competencies Table
    headers = ["Competencias Perfil de egreso", "Excelente Dominio (ED)", "Alto Dominio (AD)", "Dominio Aceptable (DA)", "Dominio Insuficiente (DP)", "Dominio no logrado (DNL)", "Comentarios"]
    
    competencies = [
        ("Crear planos de redes de datos y sistemas de telecomunicaciones según normas técnicas.", "", "X", "", "", "", "Sólida experiencia en diseño de topologías IT/OT complejas y documentación técnica bajo estándares de la industria."),
        ("Seleccionar medios de transmisión físicos, ópticos o inalámbricos adecuados.", "", "X", "", "", "", "Dominio de medios guiados y no guiados; experiencia práctica en despliegues físicos y proyectos de telecomunicaciones."),
        ("Adaptar tecnologías de punta y tendencias emergentes de conectividad.", "X", "", "", "", "", "Integración continua de Inteligencia Artificial (Gemini Live API), virtualización con Proxmox, Dokploy y orquestación moderna."),
        ("Controlar y operar redes corporativas de gran tamaño.", "X", "", "", "", "", "Manejo avanzado de enrutamiento dinámico, topologías Cisco (CCNA), segmentación de VLANs y arquitectura de alta disponibilidad."),
        ("Unificar servicios de voz, datos y video.", "X", "", "", "", "", "Implementación práctica de telefonía VoIP sobre Asterisk PBX en contenedores e integración directa con agentes de IA de voz."),
        ("Automatizar procesos y gestión de plataformas de red.", "X", "", "", "", "", "Desarrollo de scripts de automatización en Python, Google Apps Script, parsers de logs y pipelines de despliegue."),
        ("Gestionar la seguridad de la información frente a vulnerabilidades.", "X", "", "", "", "", "Especialización en ciberseguridad ofensiva/defensiva (CTF DevSec), análisis de vectores de ataque como SQLi/MQTT y hardening."),
        ("Crear planes de prevención y respuesta a riesgos informáticos.", "X", "", "", "", "", "Diseño de arquitecturas perimetrales con pfSense, Suricata en modo Inline IPS, pfBlockerNG con MaxMind GeoIP y mitigación automática."),
        ("Dirigir labores de soporte técnico avanzado para la infraestructura.", "", "X", "", "", "", "Resolución de incidencias en capa 2 a 7, troubleshooting en Linux (Arch/Ubuntu), FreeBSD/pfSense y análisis de tráfico de red."),
        ("Liderar proyectos de telecomunicaciones e integración digital.", "X", "", "", "", "", "Liderazgo técnico en iniciativas como Cyb4Students Week, 1° lugar en SummIT 5G Innovation Challenge y despliegues CTF internacionales."),
        ("Evaluar costos, calidad del servicio y cumplimiento de normativas.", "", "X", "", "", "", "Optimización de costos mediante software Open Source (pfSense, Suricata, Asterisk) y cumplimiento de estándares de seguridad.")
    ]

    t_comp = doc.add_table(rows=len(competencies)+1, cols=7)
    t_comp.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    for col_idx, text in enumerate(headers):
        cell = t_comp.rows[0].cells[col_idx]
        cell.text = text
        set_cell_background(cell, "003366")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)
                
    # Data rows
    for row_idx, data in enumerate(competencies):
        row = t_comp.rows[row_idx+1]
        for col_idx, val in enumerate(data):
            cell = row.cells[col_idx]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8.5)
                    if col_idx in [1, 2, 3, 4, 5] and val == "X":
                        run.font.bold = True
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if row_idx % 2 == 1:
            for cell in row.cells:
                set_cell_background(cell, "F9FAFC")

    doc.save(filename)
    print(f"Generated {filename}")

def create_diario_reflexion_fase1_docx(filename):
    doc = docx.Document()
    
    # Title
    title = doc.add_paragraph()
    r = title.add_run("Diario de Reflexión - Fase 1\n")
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0, 51, 102)
    
    sub = title.add_run("Experiencia de Aprendizaje 1 (EA1) | Definición Proyecto APT\nEstudiante: Bruno Urrea Ortiz | Sede San Joaquín | Duoc UC\n")
    sub.font.size = Pt(11)
    sub.font.italic = True

    sections = [
        ("1. Asignaturas y certificados de mayor motivación:",
         "¿Cuáles son las asignaturas o certificados que más motivaron tu aprendizaje y se relacionan con tus intereses profesionales?",
         "Las asignaturas que mayor impacto y motivación han generado en mi formación son aquellas orientadas a la Ciberseguridad Defensiva y Ofensiva, Seguridad Perimetral, Routing & Switching avanzado (CCNA / Cisco Networking Academy), Sistemas Operativos tipo Unix/Linux y Arquitectura de Redes.\n\n"
         "A lo largo de mi formación académica y práctica, profundicé en la concepción de que una red no solo debe ser funcional y de alta disponibilidad, sino inherentemente segura por diseño (Security by Design). Esto despertó mi pasión por la construcción de laboratorios complejos (Home Lab bare-metal con Proxmox VE y emulación con GNS3), la administración avanzada de firewalls pfSense y la detección/prevención de intrusiones con Suricata IPS, orientando mi perfil hacia la ingeniería de seguridad perimetral, la respuesta autónoma ante incidentes (SOAR) y la fortificación de infraestructuras críticas."),
        
        ("2. Análisis de competencias (Fortalezas y Oportunidades de Mejora):",
         "¿Cuáles consideras que tienes más desarrolladas y te sientes más seguro aplicando? ¿En cuáles te sientes más débil y requieren ser fortalecidas?\nSumado a lo anterior, ¿Hay alguna competencia que hayas desarrollado de forma autodidacta en alguna actividad extracurricular que quieras destacar?",
         "Fortalezas: Poseo alto dominio en la administración de entornos Linux (Arch Linux con Hyprland), virtualización bare-metal con Proxmox, diseño de redes segmentadas, securización perimetral mediante pfSense, inspección profunda de paquetes con Suricata en modo Inline IPS, y configuración de proxies inversos como HAProxy.\n\n"
         "Áreas a fortalecer: Profundizar en la estandarización y modelamiento formal de costos operacionales/financieros a nivel corporativo bajo marcos formales (ITIL/TOGAF) y pulir metodologías de comunicación ejecutiva ante comités directivos de negocio.\n\n"
         "Competencias autodidactas y extracurriculares: He desarrollado capacidades avanzadas de investigación y explotación de vulnerabilidades en protocolos industriales (ej. vectores de ataque sobre MQTT para el CTF Llaitún 2025 - Water Shield Edition), despliegue de infraestructuras CTF en vivo para eventos masivos (FIDAE 2026 con Ciberlab), y desarrollo de soluciones de automatización e integración de Inteligencia Artificial en tiempo real."),
        
        ("3. Proyección Profesional y Perfil de Egreso:",
         "¿En qué área deseas trabajar cuando egreses de tu carrera? ¿Cómo te gustaría que fuera tu escenario laboral en 5 años más? ¿Cuál es tu plan o proyecto para lograrlo? ¿Qué competencias requieres fortalecer?",
         "Área de egreso: Ingeniero de Ciberseguridad / Arquitecto de Infraestructura y SecOps, integrando ingeniería de redes perimetrales, automatización de respuesta a incidentes (SOAR) y detección de amenazas en entornos IT y OT.\n\n"
         "Escenario laboral en 5 años: Liderar un equipo de ingeniería en ciberseguridad / SOC de respuesta a incidentes o como consultor senior en arquitecturas Zero Trust y sistemas autónomos de defensa. Mi plan de acción incluye la certificación oficial Cisco CCNA 200-301, certificaciones especializadas de seguridad (eJPT/OSCP o certificaciones de arquitectura de seguridad), y la continua participación en CTFs de élite con el equipo DevSec.\n\n"
         "Competencias a fortalecer: Gestión integral de proyectos corporativos a gran escala y oratoria ejecutiva para traducir métricas técnicas de riesgo a impacto financiero y estratégico de negocio."),
        
        ("4. Definición del Proyecto APT (Portafolio de Título):",
         "Los Proyectos APT que ya habías diseñado como plan de trabajo, ¿se relacionan con tus proyecciones profesionales actuales? ¿En qué contexto se sitúa este Proyecto APT?",
         "El proyecto de título seleccionado representa la convergencia perfecta de mis competencias en redes, ciberseguridad perimetral, automatización y tecnologías emergentes:\n\n"
         "NOMBRE DEL PROYECTO: 'Sistema Autónomo de Detección IPS, Filtrado Inteligente de Falsos Positivos y Respuesta Telefónica de Incidentes en Tiempo Real mediante Agente de IA y Asterisk PBX'.\n\n"
         "DESCRIPCIÓN DEL PROYECTO:\n"
         "1. Infraestructura de Red y Perímetro: Implementación de un firewall pfSense perimetral configurado con Suricata en modo Inline IPS (Netmap), pfBlockerNG-devel alimentado con bases de datos GeoIP de MaxMind y listas de reputación de IPs maliciosas.\n"
         "2. Publicación Segura de Servicios: Despliegue de un proxy inverso HAProxy que publica hacia Internet un entorno web vulnerable controlado (DVWA) en DMZ para pruebas éticas.\n"
         "3. Motor de Correlación y Supresión de Falsos Positivos (pfctl Engine): Desarrollo de un motor en Python que inspecciona los registros EVE de Suricata y las tablas dinámicas de FreeBSD pfctl (tabla snort2c y estados de conexión), eliminando el ruido y falsos positivos habituales (más del 50% del tráfico ruidoso) para aislar ataques críticos reales (ej. SQL Injection autenticado o Command Injection).\n"
         "4. Agente de IA Conversacional en Tiempo Real y Telefonía PBX: Ante un ataque real mitigado y bloqueado por pfctl, el motor dispara un webhook/trigger hacia una centralita Asterisk PBX alojada en la nube/contenedor, la cual efectúa una llamada telefónica automática inmediata al CISO/SOC Lead. Un agente de IA (Gemini Live API) interactúa por voz con el operador, entregando un debriefing preciso del ataque (IP origen, geolocalización, vector, bloqueo confirmado) y proponiendo mitigaciones estratégicas en tiempo real.")
    ]

    for title_text, q_text, ans_text in sections:
        p_sec = doc.add_paragraph()
        r_sec = p_sec.add_run(title_text)
        r_sec.font.bold = True
        r_sec.font.size = Pt(12)
        r_sec.font.color.rgb = RGBColor(0, 51, 102)
        
        p_q = doc.add_paragraph()
        r_q = p_q.add_run(q_text)
        r_q.font.italic = True
        r_q.font.size = Pt(10)
        r_q.font.color.rgb = RGBColor(80, 80, 80)
        
        p_ans = doc.add_paragraph()
        p_ans.add_run(ans_text)
        doc.add_paragraph() # Spacer

    doc.save(filename)
    print(f"Generated {filename}")

def create_diario_reflexion_fase2_docx(filename):
    doc = docx.Document()
    
    # Title
    title = doc.add_paragraph()
    r = title.add_run("Diario de Reflexión - Fase 2\n")
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0, 51, 102)
    
    sub = title.add_run("Experiencia de Aprendizaje 2 (EA2) | Desarrollo y Monitoreo del Proyecto APT\nEstudiante: Bruno Urrea Ortiz | Sede San Joaquín | Duoc UC\n")
    sub.font.size = Pt(11)
    sub.font.italic = True

    sections = [
        ("1. Seguimiento de Carta Gantt y Cumplimiento de Cronograma:",
         "¿Has podido cumplir todas las actividades en los tiempos definidos? ¿Qué factores han facilitado o dificultado el desarrollo de las actividades de tu plan de trabajo?",
         "El cronograma establecido en la Carta Gantt se ha cumplido rigurosamente y dentro de los plazos proyectados.\n\n"
         "Factores facilitadores: La amplia experiencia previa en virtualización con Proxmox VE y configuración de pfSense aceleró significativamente el despliegue del entorno base de laboratorio, la integración de HAProxy y el setup de Suricata en modo Inline IPS.\n\n"
         "Factores de dificultad: El principal desafío técnico consistió en afinar la correlación de eventos entre el archivo eve.json de Suricata y las tablas de estados de FreeBSD pfctl, debido a la alta tasa de ruido y firmas genéricas que generan falsas alertas en entornos de prueba web."),
        
        ("2. Resolución de Dificultades y Contingencias Técnicas:",
         "¿De qué manera has enfrentado y/o planeas enfrentar las dificultades que han afectado el desarrollo de tu Proyecto APT?",
         "Para solucionar la tasa de falsos positivos característica de los motores IDS/IPS tradicionales, se diseñó un algoritmo de doble verificación en el 'pfctl Log Engine':\n"
         "1. Verificación de payload HTTP y decodificación de parámetros SQLi en la capa de aplicación expuesta por HAProxy.\n"
         "2. Consulta en tiempo real a la tabla snort2c de pfctl ('pfctl -t snort2c -T show') para confirmar que la dirección IP atacante fue efectivamente expulsada y bloqueada en la capa de red del kernel FreeBSD.\n\n"
         "Adicionalmente, se configuró un pipeline asíncrono con WebSockets para conectar el motor de eventos con la API de voz Gemini Live, asegurando baja latencia en la llamada iniciada por Asterisk PBX."),
        
        ("3. Evaluación de Evidencias de Avance:",
         "¿Cómo evalúas tu(s) evidencia(s) de avance? ¿Qué destacas y qué podrías hacer para mejorar tus evidencias?",
         "Las evidencias acumuladas son de alto estándar técnico y 100% verificables:\n"
         "- Capturas y logs de bloqueo activo en pfSense (Suricata + pfctl).\n"
         "- Grabaciones de audio y logs de telemetría de las llamadas ejecutadas por Asterisk PBX y el agente de IA Gemini Live.\n"
         "- Repositorio estructurado en GitHub con código modular, contenedores Docker y configuraciones reproducibles.\n\n"
         "Oportunidad de mejora: Incorporar un panel de métricas visuales (dashboard interactivo) que grafique el tiempo de respuesta total desde la detección del paquete malicioso hasta el descolgado telefónico de la llamada de alerta."),
        
        ("4. Inquietudes y Consultas para el Docente / Pares:",
         "¿Qué inquietudes te quedan sobre cómo proceder? ¿Qué pregunta te gustaría hacerle a tu docente o a tus pares?",
         "Consulta técnica y pedagógica: ¿Cuál es el criterio preferido por la comisión evaluadora para la demostración en vivo de la llamada telefónica del agente de IA durante la defensa de la Fase 3? ¿Es recomendable proyectar el flujo de paquetes en Wireshark/SIP en paralelo al audio en directo de la llamada con Gemini Live para maximizar el impacto de la rúbrica?"),
        
        ("5. Gestión y Distribución del Trabajo en Equipo:",
         "¿Consideran que las actividades deben ser redistribuidas entre los miembros del grupo? ¿Hay nuevas actividades que deban ser asignadas a algún miembro del grupo?",
         "La distribución de tareas se encuentra equilibrada y alineada con las fortalezas individuales de cada integrante del equipo (infraestructura de red, configuración de servicios web, telefonía IP y documentación académica). En esta fase se acordó asignar la preparación del guión de prueba de inyecciones SQL simuladas y la recolección de evidencias para el informe final."),
        
        ("6. Evaluación del Trabajo Grupal:",
         "¿Cómo evalúan el trabajo en grupo? ¿Qué aspectos positivos destacan? ¿Qué aspectos podrían mejorar?",
         "Evaluación positiva y colaborativa. Destaco el compromiso constante, la comunicación fluida a través de canales técnicos y la capacidad de articular distintas áreas de la carrera (Routing, Switching, Telefonía, Ciberseguridad y Programación). Como punto a optimizar, debemos sincronizar con mayor frecuencia las versiones de los documentos en el repositorio para evitar desfases previos a las entregas de hito.")
    ]

    for title_text, q_text, ans_text in sections:
        p_sec = doc.add_paragraph()
        r_sec = p_sec.add_run(title_text)
        r_sec.font.bold = True
        r_sec.font.size = Pt(12)
        r_sec.font.color.rgb = RGBColor(0, 51, 102)
        
        p_q = doc.add_paragraph()
        r_q = p_q.add_run(q_text)
        r_q.font.italic = True
        r_q.font.size = Pt(10)
        r_q.font.color.rgb = RGBColor(80, 80, 80)
        
        p_ans = doc.add_paragraph()
        p_ans.add_run(ans_text)
        doc.add_paragraph() # Spacer

    doc.save(filename)
    print(f"Generated {filename}")

if __name__ == "__main__":
    create_autoevaluacion_docx("Urrea_Bruno_1.1_APT122_AutoevaluacionCompetenciasFase1.docx")
    create_diario_reflexion_fase1_docx("Urrea_Bruno_1.2_APT122_DiarioReflexionFase1.docx")
    create_diario_reflexion_fase2_docx("Urrea_Bruno_2.1_APT122_DiarioReflexionFase2.docx")
